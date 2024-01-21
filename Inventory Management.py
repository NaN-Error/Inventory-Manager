"""
Inventory Management System

This program provides a comprehensive solution for managing and tracking inventory. It is built with a graphical user interface (GUI) using tkinter, making it user-friendly and accessible. The key features include:

- Database Management: Features a `DatabaseManager` class that handles all database operations using sqlite3. It includes functions to setup the database and manage folder paths, ensuring organized data storage and retrieval.
- GUI Development: Utilizes tkinter for creating interactive windows, dialogs, and widgets, facilitating user interaction for various inventory tasks.
- File Handling: Supports operations with different file formats (e.g., Excel, Word) using libraries like pandas, openpyxl, and python-docx, enabling users to import and export inventory data.
- Date and Time Functions: Incorporates datetime and dateutil modules for managing dates, essential for tracking inventory timelines.
- Regular Expressions: Uses the re module for text processing, ensuring data validation and formatting.
- Image Processing: Implements PIL and openpyxl_image_loader for handling and displaying images, useful in visual inventory management.
- Multithreading: Applies threading to enhance performance and responsiveness of the application.
- Logging: Includes logging functionality to track and record application activities, aiding in debugging and maintenance.

This system is designed to be robust and versatile, suitable for various inventory management needs. It is especially beneficial for small to medium-sized businesses or personal inventory tracking.

Dependencies:
- tkinter, pandas, sqlite3, openpyxl, python-docx, Pillow, tkcalendar, ttkthemes, and additional Python standard libraries.

Usage:
Execute the script to start the inventory management application. Use the GUI to perform operations like adding, updating, deleting, and viewing inventory items.

Author: [WB]
Version: [Your Version]

Note: The docstring provides a general overview. Detailed documentation for each class and function within the code is recommended for better understanding and maintenance.
"""

import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
import pandas as pd
from docx import Document
import sqlite3
from tkinter import END
from tkinter import Toplevel
from openpyxl import load_workbook
import re
import subprocess
import sys
import openpyxl
import webbrowser
from pathlib import Path
from tkcalendar import Calendar
from tkinter.font import Font
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.styles import Alignment
from ttkthemes import ThemedTk
from openpyxl import Workbook
import math
from decimal import Decimal, ROUND_HALF_UP
from decimal import Decimal, InvalidOperation
from openpyxl.styles import PatternFill
from io import BytesIO
import threading
import io
from tkinter import simpledialog
from PIL import Image, ImageTk
from openpyxl_image_loader import SheetImageLoader
from tkinter import Label, Toplevel
import logging
from logging.handlers import RotatingFileHandler
import time
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt


# Prototyping (make it work, then make it pretty.)
# change Load workbook to dataframe on load. (speed optimization)


class DatabaseManager: #DB practice(use txt/json to store folder paths when program finished for faster reads.)

    def __init__(self, db_name='inventory_management.db'):
        self.conn = sqlite3.connect(db_name)
        self.cur = self.conn.cursor()
        self.setup_database()

    def setup_database(self):
        self.cur.execute('''
            CREATE TABLE IF NOT EXISTS folder_paths (
                Folder TEXT PRIMARY KEY,
                Path TEXT
            )
        ''')
        self.conn.commit()

    def save_folder_path(self, folder, path):
        self.cur.execute('''
            INSERT INTO folder_paths (Folder, Path) VALUES (?, ?)
            ON CONFLICT(Folder) DO UPDATE SET Path = excluded.Path;
        ''', (folder, path))
        self.conn.commit()

    def delete_folder_path(self, old_folder_name):
        """
        Deletes the folder path record with the given old_folder_name from the database.
        """
        try:
            self.cur.execute('DELETE FROM folder_paths WHERE Folder = ?', (old_folder_name,))
            self.conn.commit()
        except Exception as e:
            print(f"Error deleting folder {old_folder_name}: {e}")

    def get_folder_path(self, folder_name):
        self.cur.execute('SELECT Path FROM folder_paths WHERE Folder = ?', (folder_name,))
        result = self.cur.fetchone()
        return result[0] if result else None

    def get_all_folders(self):
        self.cur.execute('SELECT Folder FROM folder_paths')
        return [row[0] for row in self.cur.fetchall()]

    def delete_all_folders(self):
        self.cur.execute('DELETE FROM folder_paths')
        self.conn.commit()
        
    def commit_changes(self):
        self.conn.commit()
        
    def __del__(self):
        if hasattr(self, 'conn'):
            self.conn.close()

class ExcelManager:

    def __init__(self, filepath=None, sheet_name=None):
        self.filepath = filepath
        self.sheet_name = sheet_name
        self.data_frame = None

    def load_data(self):
        if self.filepath and self.sheet_name:
            self.data_frame = pd.read_excel(self.filepath, sheet_name=self.sheet_name, engine='openpyxl')
            # Cast all columns to object dtype after loading data
            self.data_frame = self.data_frame.astype('object')

    def get_product_info(self, product_id):
        if self.data_frame is not None:
            # Convert both the product_id and the 'Product ID' column to upper case for comparison
            query_result = self.data_frame[self.data_frame['Product ID'].str.upper() == product_id.upper()]
            if not query_result.empty:
                return query_result.iloc[0].to_dict()
        return None

    def save_product_info(self, product_id, product_data):
        if self.filepath:
            try:
                #print(f"Loading workbook from {self.filepath}")
                workbook = load_workbook(self.filepath)
                #print(f"Accessing sheet {self.sheet_name}")
                sheet = workbook[self.sheet_name]

                # Start by finding the column index for product IDs
                product_id_col_index = self.get_column_index_by_header(sheet, 'Product ID')
                if not product_id_col_index:
                    #print("Product ID column not found")
                    return

                # Update product_data dictionary to convert boolean to YES/NO strings
                for key, value in product_data.items():
                    if isinstance(value, bool):
                        product_data[key] = 'YES' if value else 'NO'

                # Now iterate over the rows to find the matching product ID
                for row in sheet.iter_rows(min_col=product_id_col_index, max_col=product_id_col_index):
                    cell = row[0]
                    if cell.value and str(cell.value).strip().upper() == product_id.upper():
                        row_num = cell.row
                        for key, value in product_data.items():
                            col_index = self.get_column_index_by_header(sheet, key)
                            if col_index:
                                # Special handling for 'To Sell After' date
                                if key == 'To Sell After' and isinstance(value, datetime):
                                    value = value.strftime('%m/%d/%Y')  # Format the date
                                    sheet.cell(row=row_num, column=col_index, value=value)
                                elif key == 'Fair Market Value':
                                    # Convert value to float if it's not None or empty
                                    value = float(value) if value else 0
                                    # Set the cell value
                                    cell = sheet.cell(row=row_num, column=col_index, value=value)
                                    # Set the number format for currency
                                    cell.number_format = '"$"#,##0.00'
                                else:
                                    sheet.cell(row=row_num, column=col_index, value=value)
                        workbook.save(self.filepath)
                        break
                else:
                    #print(f"Product ID {product_id} not found in the sheet.")
                    pass
            except Exception as e:
                #print(f"Failed to save changes to Excel file: {e}")
                raise

    @staticmethod
    def get_column_index_by_header(sheet, header_name):
        """
        Gets the column index based on the header name.
        :param sheet: The sheet to search in.
        :param header_name: The header name to find.
        :return: The index of the column, or None if not found.
        """
        for col in sheet.iter_rows(min_row=1, max_row=1, values_only=True):
            if header_name in col:
                return col.index(header_name) + 1
        return None

class Application(tk.Frame):

    def __init__(self, master=None):
        super().__init__(master)
        self.db_manager = DatabaseManager()
        self.excel_manager = ExcelManager()
        self.edit_mode = False  # Add this line to initialize the edit_mode attribute
        self.inventory_folder = None
        self.sold_folder = None
        self.to_sell_folder = None
        self.pack(fill='both', expand=True)
        self.last_changed = None
        self.initial_discount_price = None  # Class attribute to store the initial discount price
        self.initial_percent_discount = None  # Class attribute to store the initial discount price
        self.initial_product_price_plus_ivu = ''  # Initialize the variable
        self.trigger_price_focus_out_flag = True
        self.running = True
        self.current_product_id = None
        self.workbook_cache = None
        self.workbook_path = None
        self.image_cache = {}
        #self.trigger_save_flag = False # Can be used to save when pressing enter once while in Product Price (+IVU) entry.

        self.configure_logger()
        self.cache_images_on_load()
        self.load_settings()
        self.Main_Window_Widgets() 
        self.combine_and_display_folders()
        self.master.update_idletasks()
        self.update_excel_file_on_start_question()
        #self.first_run()
        #remove update_folders_path function?

    def configure_logger(self):
        # Set up a logger
        self.logger = logging.getLogger('InventoryManagementLogger')
        self.logger.setLevel(logging.INFO)  # Set the logging level

        # Create a rotating file handler
        handler = RotatingFileHandler('inventory_management.log', maxBytes=10000000, backupCount=5, encoding='utf-8')
        handler.setLevel(logging.INFO)

        # Create a logging format
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)

        # Add the handler to the logger
        self.logger.addHandler(handler)

        # Log the start of the application
        self.logger.info("----Inventory Management Application started----")

    def cache_images_on_load(self):
        self.logger.info("Starting to cache images on load")

        # Load Excel settings
        filepath, sheet_name = self.load_excel_path_and_sheet()
        if filepath and sheet_name:
            # Validate file and sheet
            if not os.path.exists(filepath):
                self.logger.error(f"Excel file not found at {filepath}. Skipping image caching.")
                return
            # Additional validation for sheet can be added here if necessary

            self.logger.info(f"Excel settings loaded with filepath: {filepath} and sheet_name: {sheet_name}")
            try:
                self.cache_images(filepath, sheet_name)
                self.logger.info("Images have been successfully cached")
            except Exception as e:
                self.logger.error(f"An error occurred while caching images: {e}")
        else:
            self.logger.error("Failed to load Excel settings or they are incomplete. Skipping image caching.")

    def load_settings(self):
        self.logger.info("Attempting to load folders' paths from file")
        try:
            with open("folders_paths.txt", "r") as file:
                lines = file.read().splitlines()
                self.inventory_folder = lines[0]
                self.sold_folder = lines[1]
                self.to_sell_folder = lines[2] if len(lines) > 2 else None

                self.logger.info(f"Loaded paths: Inventory - {self.inventory_folder}, Sold - {self.sold_folder}, To Sell - {self.to_sell_folder}")
        except FileNotFoundError:
            self.logger.error("folders_paths.txt not found. Paths not loaded.")

    # def save_settings(self):
    #     self.logger.info("Saving settings for inventory and sold folders")

    #     try:
    #         # Update the table with the new paths for the inventory folder
    #         self.db_manager.cur.execute('''
    #             UPDATE folder_paths SET Path = ? WHERE Folder = 'Root Folder'
    #         ''', (self.inventory_folder,))
    #         self.logger.info(f"Inventory folder path updated to: {self.inventory_folder}")

    #         # Update the table with the new paths for the sold folder
    #         self.db_manager.cur.execute('''
    #             UPDATE folder_paths SET Path = ? WHERE Folder = 'Sold'
    #         ''', (self.sold_folder,))
    #         self.logger.info(f"Sold folder path updated to: {self.sold_folder}")

    #         # Commit the changes to the database
    #         self.db_manager.conn.commit()
    #         self.logger.info("Settings saved successfully")

    #     except Exception as e:
    #         self.logger.error(f"Error saving settings: {e}")
    
    def update_excel_file_on_start_question(self):
        """
        Displays a dialog asking the user if they want to update Excel empty fields.
        Calls first_run if the user responds affirmatively.
        """
        root = tk.Tk()
        root.withdraw()
        self.logger.info("Asking user to update Excel data.")
        user_response = messagebox.askyesno("Update Excel Data", "Do you want to update the Excel empty fields?")
        root.destroy()

        if user_response:
            self.logger.info("User chose to update Excel data.")
            self.first_run()
        else:
            self.logger.info("User chose not to update Excel data.")
            pass

    def first_run(self):
        """
        Executes a series of operations including updating Excel data, updating prices,
        updating folder paths, generating a report of products to sell, and checking for missing Word documents.
        """
        self.logger.info("Starting first run operations.")
        self.update_excel_data()
        self.update_prices()
        self.update_all_folder_paths_and_names()
        self.products_to_sell_report()
        self.check_for_missing_word_docs()
        self.logger.info("Completed first run operations.")


    def Main_Window_Widgets(self):
        self.logger.info("Initializing main window widgets")
        try:
            self.top_frame = ttk.Frame(self)
            self.top_frame.pack(fill='x')

            self.settings_button = ttk.Button(self.top_frame, text='Settings', command=self.Settings_Window_Start)
            self.settings_button.pack(side='right')

            self.search_frame = ttk.Frame(self)
            self.search_frame.pack(fill='x')

            self.search_label = ttk.Label(self.search_frame, text="Enter product name here:")
            self.search_label.pack(anchor='w')

            self.search_entry = ttk.Entry(self.search_frame, width=30)  # Same width as the Listbox
            self.search_entry.pack(side='left', fill='x', anchor='w')
            self.search_entry.bind('<KeyRelease>', self.search)

            self.bottom_frame = ttk.Frame(self)
            self.bottom_frame.pack(fill='both', expand=True)

            self.list_outer_frame = ttk.Frame(self.bottom_frame)
            self.list_outer_frame.pack(side='left', fill='y')

            self.list_frame = ttk.Frame(self.list_outer_frame)
            self.list_frame.pack(side='left', fill='both', expand=True)

            self.folder_list = tk.Listbox(self.list_frame, width=30)
            self.folder_list.pack(side='left', fill='both', expand=False)
            self.folder_list.bind('<<ListboxSelect>>', self.display_product_details)

            self.folder_list.bind('<Down>', self.next_product)
            self.folder_list.bind('<Up>', self.previous_product)

            self.list_scrollbar = ttk.Scrollbar(self.list_frame)
            self.list_scrollbar.pack(side='right', fill='y')
            self.folder_list.config(yscrollcommand=self.list_scrollbar.set)
            self.list_scrollbar.config(command=self.folder_list.yview)
            
            self.logger.info("Main window widgets initiated")
            self.Product_Form()        

        except Exception as e:
            self.logger.error(f"Error initializing main window widgets: {e}")

    def next_product(self, event):
        if self.folder_list.size() > 0:
            current_selection = self.folder_list.curselection()
            if current_selection:
                next_index = current_selection[0] + 1
                if next_index < self.folder_list.size():
                    self.folder_list.selection_clear(current_selection)
                    self.folder_list.selection_set(next_index)
                    self.folder_list.see(next_index)

    def previous_product(self, event):
        if self.folder_list.size() > 0:
            current_selection = self.folder_list.curselection()
            if current_selection:
                prev_index = current_selection[0] - 1
                if prev_index >= 0:
                    self.folder_list.selection_clear(current_selection)
                    self.folder_list.selection_set(prev_index)
                    self.folder_list.see(prev_index)

    def combine_and_display_folders(self):
        """
        Combines and displays the folder names from various paths including inventory, sold, 
        to sell, damaged, and personal folders. Updates these folder paths in the database. 
        The folder list is first cleared, then updated with the combined and sorted folder names.
        """
        self.logger.info("Combining and displaying folders")

        # Clear the folder list first
        self.folder_list.delete(0, tk.END)

        # Initialize additional folders based on the inventory folder
        if self.inventory_folder:
            parent_dir = os.path.dirname(self.inventory_folder)
            self.damaged_folder = os.path.join(parent_dir, "Damaged")
            self.personal_folder = os.path.join(parent_dir, "Personal")

            # Create additional folders if they don't exist
            for folder in [self.damaged_folder, self.personal_folder]:
                if not os.path.exists(folder):
                    os.makedirs(folder)

        # Begin a transaction
        self.db_manager.cur.execute("BEGIN")
        try:
            # Combine the folders from all paths including damaged and personal folders
            combined_folders = []
            for folder_path in [self.inventory_folder, self.sold_folder, self.to_sell_folder, self.damaged_folder, self.personal_folder]:
                if folder_path and os.path.exists(folder_path):
                    for root, dirs, files in os.walk(folder_path):
                        for dir_name in dirs:
                            combined_folders.append(dir_name)
                            full_path = os.path.join(root, dir_name)
                            # Update the database with the current folder paths
                            self.db_manager.cur.execute("INSERT OR REPLACE INTO folder_paths (Folder, Path) VALUES (?, ?)", (dir_name, full_path))
            self.db_manager.conn.commit()  # Commit the transaction if all is well
        except Exception as e:
            self.db_manager.conn.rollback()  # Rollback if there was an error
            self.logger.error(f"Database error in combine_and_display_folders: {e}")

        # Deduplicate folder names
        unique_folders = list(set(combined_folders))

        # Sort using the custom sort key function
        sorted_folders = sorted(unique_folders, key=self.custom_sort_key)

        # Insert the sorted folders into the list widget
        for folder in sorted_folders:
            self.folder_list.insert(tk.END, folder)
        self.logger.info("Folders combined, sorted, and displayed")

    def search(self, event):
        """
        Searches for folders based on the user's input in the search entry. 
        The search is case-insensitive and looks for matches in all relevant folders including
        inventory, sold, to sell, damaged, and personal folders.
        """
        self.logger.info("Performing search based on user input")
        search_terms = self.search_entry.get().split()  # Split the search string into words
        if search_terms:
            self.folder_list.delete(0, tk.END)  # Clear the current list

            # Define a list of folder paths to search in
            search_paths = [
                self.inventory_folder,
                self.sold_folder,
                self.to_sell_folder,
                self.damaged_folder,
                self.personal_folder
            ]

            # Filter out None or invalid paths
            valid_search_paths = [path for path in search_paths if path and os.path.exists(path)]

            # Create a list to store matching folder names
            matching_folders = []

            # Perform the search in each valid path
            for path in valid_search_paths:
                for root, dirs, files in os.walk(path):
                    # Check if 'dirs' is empty, meaning 'root' is a leaf directory
                    if not dirs:
                        folder_name = os.path.basename(root)  # Get the name of the leaf directory
                        # Check if all search terms are in the folder name (case insensitive)
                        if all(term.upper() in folder_name.upper() for term in search_terms):
                            matching_folders.append(folder_name)

            # Sort the matching folder names alphabetically
            matching_folders.sort()

            # Insert the sorted folder names into the list widget
            for folder_name in matching_folders:
                self.folder_list.insert(tk.END, folder_name)

            self.logger.info("Search completed and sorted results displayed")

        else:
            self.combine_and_display_folders()  # If the search box is empty, display all folders   
            self.logger.info("Search box is empty, displaying all folders")


# Settings Window with functions used in it.
    def Settings_Window_Start(self):
        """
        Opens the settings window where the user can configure various options like 
        choosing inventory, sold inventory, and products to sell folders, selecting the Excel database, 
        and other settings related to product data and reports.
        """
        self.logger.info("Opening settings window")

        if hasattr(self, 'settings_window') and self.settings_window.winfo_exists():
            self.settings_window.lift()
            return
        self.settings_window = tk.Toplevel(self)
        self.settings_window.title("Settings")
        self.settings_window.state('zoomed')

        # Create and grid the settings frame
        self.settings_frame = tk.Frame(self.settings_window)
        self.settings_frame.grid(row=1, column=1, sticky='nw')

        # Load settings
        self.default_filepath, self.default_sheet = self.load_excel_path_and_sheet()
        
        # Configure the grid columns of the frame
        self.settings_frame.grid_columnconfigure(1, weight=1)  # Adjust the weight as needed

        # Now grid all widgets onto the settings_frame
        self.inventory_folder_button = ttk.Button(self.settings_frame, text="Choose Inventory Folder", command=self.choose_inventory_folder)
        self.inventory_folder_button.grid(row=1, column=0, padx=5, pady=5, sticky='w')
        self.inventory_folder_label = ttk.Label(self.settings_frame, text=self.inventory_folder if self.inventory_folder else "Not chosen")
        self.inventory_folder_label.grid(row=1, column=1, padx=5, pady=5, sticky='w')

        self.sold_folder_button = ttk.Button(self.settings_frame, text="Choose Sold Inventory Folder", command=self.choose_sold_folder)
        self.sold_folder_button.grid(row=2, column=0, padx=5, pady=5, sticky='w')
        self.sold_folder_label = ttk.Label(self.settings_frame, text=self.sold_folder if self.sold_folder else "Not chosen")
        self.sold_folder_label.grid(row=2, column=1, padx=5, pady=5, sticky='w')

        self.to_sell_folder_button = ttk.Button(self.settings_frame, text="Choose Products to Sell Folder", command=self.choose_to_sell_folder)
        self.to_sell_folder_button.grid(row=3, column=0, padx=5, pady=5, sticky='w')
        self.to_sell_folder_label = ttk.Label(self.settings_frame, text=self.to_sell_folder if self.to_sell_folder else "Not chosen")
        self.to_sell_folder_label.grid(row=3, column=1, padx=5, pady=5, sticky='w')

        self.excel_db_button = ttk.Button(self.settings_frame, text="Select Excel Database", command=self.select_excel_database)
        self.excel_db_button.grid(row=4, column=0, padx=5, pady=5, sticky='w')
        excel_db_text = f"{self.default_filepath} - Sheet: {self.default_sheet}" if self.default_filepath and self.default_sheet else "Not chosen"
        self.excel_db_label = ttk.Label(self.settings_frame, text=excel_db_text)
        self.excel_db_label.grid(row=4, column=1, padx=5, pady=5, sticky='w')

        self.create_word_files_button = ttk.Button(self.settings_frame, text="Create Word Files for Products", command=self.check_for_missing_word_docs)
        self.create_word_files_button.grid(row=5, column=0, padx=5, pady=5, sticky='w')

        self.autofill_links_asin_tosellafter_data_button = ttk.Button(self.settings_frame, text="Autofill Excel Data(link, asin, tosellafter)", command=self.update_excel_data)
        self.autofill_links_asin_tosellafter_data_button.grid(row=6, column=0, padx=5, pady=5, sticky='w')

        self.update_foldersnames_folderpaths_button = ttk.Button(self.settings_frame, text="Update folder names and paths", command=self.update_all_folder_paths_and_names)
        self.update_foldersnames_folderpaths_button.grid(row=7, column=0, padx=5, pady=5, sticky='w')

        self.products_to_sell_list_button = ttk.Button(self.settings_frame, text="Show list of products available to sell", command=self.products_to_sell_report)
        self.products_to_sell_list_button.grid(row=8, column=0, padx=5, pady=5, sticky='w')

        self.update_prices_button = ttk.Button(self.settings_frame, text="Update empty product prices based on Fair Market Value.", command=self.update_prices)
        self.update_prices_button.grid(row=9, column=0, padx=5, pady=5, sticky='w')

        self.update_prices_button = ttk.Button(self.settings_frame, text="First run.", command=self.first_run)
        self.update_prices_button.grid(row=10, column=0, padx=5, pady=5, sticky='w')

        self.back_button = ttk.Button(self.settings_window, text="<- Back", command=self.back_to_main)
        self.back_button.grid(row=0, column=0, sticky='w', padx=5, pady=5)


        self.logger.info("Settings window initialized and displayed")

        self.combine_and_display_folders()
        self.settings_window.protocol("WM_DELETE_WINDOW", lambda: on_close(self, self.master))
        self.master.withdraw()
    
    def create_all_word_docs(self):
        """
        Creates Word documents for all items listed in the correlation window.
        """

        # Log the start of creating all Word documents
        self.logger.info("Starting the creation of all Word documents")

        for iid in self.correlate_tree.get_children():
            item_values = self.correlate_tree.item(iid, 'values')
            doc_data = (item_values[0], item_values[1], item_values[2])

            # Log the data of each item being processed
            self.logger.info(f"Creating Word document for item: {doc_data}")

            self.create_word_doc(doc_data, iid, show_message=False)

        messagebox.showinfo("Success", "All Word documents have been created.")
        self.correlate_window.destroy()
        self.Settings_Window_Start()

        # Log the completion of creating all Word documents
        self.logger.info("All Word documents created successfully")

    def create_word_doc(self, doc_data, iid, show_message=True):
        """
        Creates a Word document for a specific product, pulling relevant information from the Excel data. 
        The document includes details like product ID, name, price, link, and comments.
        """
        def safe_format_currency(value):
            try:
                return f"${float(value):.2f}" if value is not None else "N/A"
            except ValueError:
                return str(value)

        def safe_format_percentage(value):
            try:
                return f"{float(value)}%" if value is not None else "N/A"
            except ValueError:
                return str(value)
            
        def add_styled_paragraph(doc, text, variable_text):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.italic = True
            run.underline = True
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN  # Applying light green highlight
            run.font.size = Pt(12)  # Setting font size to 12
            p.add_run(variable_text)


        # Log the start of the Word document creation process
        self.logger.info(f"Creating Word document for product ID {doc_data[1]}")

        # Unpack the data tuple
        folder_name, product_id, product_name = doc_data
        # Retrieve the folder path from the database
        folder_path = self.get_folder_path_from_db(str(product_id))

        if folder_path:
            try:
                # Retrieve the product link
                product_price_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Product Price']
                if not product_price_series.empty:
                    product_price = product_price_series.iloc[0]
                else:
                    product_price = "N/A"  # Default to "N/A" if not found

                # Retrieve the product link
                ivu_tax_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'IVU Tax']
                if not ivu_tax_series.empty:
                    ivu_tax = ivu_tax_series.iloc[0]
                else:
                    ivu_tax = "N/A"  # Default to "N/A" if not found

                # Retrieve the product link
                product_price_after_ivu_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Product Price After IVU']
                if not product_price_after_ivu_series.empty:
                    product_price_after_ivu = product_price_after_ivu_series.iloc[0]
                else:
                    product_price_after_ivu = "N/A"  # Default to "N/A" if not found

                # Retrieve the product link
                order_link_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Order Link']
                if not order_link_series.empty:
                    order_link = order_link_series.iloc[0]
                else:
                    order_link = "N/A"  # Default to "N/A"            
                    
                # Retrieve the product description
                product_description_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Product Description']
                if not product_description_series.empty and not pd.isna(product_description_series.iloc[0]):
                    product_description = product_description_series.iloc[0]
                else:
                    product_description = "No Product Description At The Moment"

                # Retrieve the comments
                comments_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Comments']
                if not comments_series.empty and not pd.isna(comments_series.iloc[0]):
                    comments = comments_series.iloc[0]
                else:
                    comments = "No Comments Found"
                    
                    # Retrieve the product 
                product_name_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Product Name']
                if not product_name_series.empty:
                    product_name = product_name_series.iloc[0]
                else:
                    product_name = "N/A"  # Default to "N/A" 

                    # Retrieve the product 
                discount_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Discount']
                if not discount_series.empty:
                    discount = discount_series.iloc[0]
                else:
                    discount = "N/A"  # Default to "N/A" 

                    # Retrieve the product 
                discount_percentage_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Discount Percentage']
                if not discount_percentage_series.empty:
                    discount_percentage = discount_percentage_series.iloc[0]
                else:
                    discount_percentage = "N/A"  # Default to "N/A" 

            except Exception as e:
                self.logger.info(f"Error retrieving data: {e}")  # Debugging print statement

            # Path for the new Word document named 'Product Information.docx'
            doc_path = os.path.join(folder_path, 'Product Information.docx')
            try:
                # Create a new Word document
                doc = Document()

                # Convert all values to strings with appropriate formatting
                product_id_str = str(product_id)
                product_name_str = str(product_name) if product_name is not None else "N/A"
                product_price_str = safe_format_currency(product_price)
                ivu_tax_str = safe_format_currency(ivu_tax)
                product_price_after_ivu_str = safe_format_currency(product_price_after_ivu)
                discount_str = safe_format_currency(discount)
                discount_percentage_str = safe_format_percentage(discount_percentage)
                product_description_str = str(product_description) if product_description is not None else "N/A"
                comments_str = str(comments) if comments is not None else "N/A"

                # Adding styled paragraphs with specified font size
                add_styled_paragraph(doc, "Product ID: ", product_id_str)
                add_styled_paragraph(doc, "Product Name: ", product_name_str)
                doc.add_paragraph("")  # Empty line
                add_styled_paragraph(doc, "Product Price: ", product_price_str)
                add_styled_paragraph(doc, "IVU Tax: ", ivu_tax_str)
                add_styled_paragraph(doc, "Product Price After IVU (Sale Price): ", product_price_after_ivu_str)
                add_styled_paragraph(doc, "Reseller Earnings : ", f"{discount_str}     [ = {discount_percentage_str} of {product_price_str} (Product Price)]")
                doc.add_paragraph("")  # Empty line
                add_styled_paragraph(doc, "Product Description:", "")
                doc.add_paragraph(product_description_str)
                doc.add_paragraph("")  # Empty line
                add_styled_paragraph(doc, "Comments:", "")
                doc.add_paragraph(comments_str)

                # Uncomment the next line if you want to include the order link
                # add_styled_paragraph(doc, "Amazon Link (to get the product description and pictures, if needed): ", order_link)


                # Save the document
                doc.save(doc_path)

                if show_message:
                    messagebox.showinfo("Document Created", f"Word document for '{product_id}' has been created successfully.")
                    self.logger.info(f"Word document for product ID {product_id} created successfully")

                # Check if 'correlate_tree' exists before trying to delete an item
                if hasattr(self, 'correlate_tree'):
                    try:
                        self.correlate_tree.delete(iid)
                    except Exception as e:
                        self.logger.error(f"Error while updating the Treeview: {e}")
                # Bring the correlate_window back to the top

                if hasattr(self, 'correlate_window'):
                    self.correlate_window.lift()

                if hasattr(self, 'correlate_tree') and not self.correlate_tree.get_children():
                    self.correlate_window.destroy()
                    self.Settings_Window_Start()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to create document for Product ID {product_id}: {e}")
                self.logger.error(f"Failed to create document for product ID {product_id}: {e}")
        else:
            messagebox.showerror("Error", f"No folder found for Product ID {product_id}")
            self.logger.error(f"No folder found for product ID {product_id}")

    def check_for_missing_word_docs(self):
        """
        Correlates data between the Excel file and Word documents. 
        It checks if each product in the Excel file has an associated Word document in its respective folder.
        Notifies if any Word documents are missing for the products listed in the Excel file.
        """

        # Log when correlation process is initiated
        self.logger.info("Initiating data correlation between Excel file and Word documents")
        
        filepath, sheet_name = self.load_excel_path_and_sheet()

        # Check if the Excel settings are properly loaded
        if not filepath or not sheet_name:
            messagebox.showerror("Error", "Excel database settings not found.")
            self.logger.error("Excel database settings not found, correlation aborted")
            return

        # Load the data into the ExcelManager instance
        self.excel_manager.filepath = filepath  # Set the filepath
        self.excel_manager.sheet_name = sheet_name  # Set the sheet name
        self.excel_manager.load_data()  # Load the data
        
        try:
            # Load Excel data
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            self.logger.info("Excel data loaded successfully")
            product_ids = df['Product ID'].tolist()
            #print(f"Product IDs from Excel: {product_ids}")
        except Exception as e:
            messagebox.showerror("Error", f"Unable to load Excel file: {str(e)}")
            self.logger.error(f"Unable to load Excel file: {e}")

            return
            # Filter out nan values from the product_ids list using pandas notnull function
            
        # Sort the DataFrame based on 'Product ID'
        df_sorted = df.sort_values('Product ID').dropna(subset=['Product ID'])

        # Filter out nan values from the product_ids list
        product_ids = df_sorted['Product ID'].tolist()
        #print(f"Sorted and Filtered Product IDs from Excel: {product_ids}")
        
        missing_docs = []
        for product_id in product_ids:
            folder_path = self.get_folder_path_from_db(str(product_id))
            if folder_path:
                # Check specifically for 'Product Information.docx' file
                if not os.path.isfile(os.path.join(folder_path, 'Product Information.docx')):
                    product_name = df.loc[df['Product ID'] == product_id, 'Product Name'].iloc[0]
                    missing_docs.append((os.path.basename(folder_path), product_id, product_name))


        #print(f"Missing documents: {missing_docs}")
        if missing_docs:
            self.prompt_missing_word_docs(missing_docs)            
            self.logger.info("Missing Word documents found, prompting user for action")
        else:
            messagebox.showinfo("Check complete", "No missing Word documents found.")
            self.logger.info("No missing Word documents found, check complete")
        # Filter out nan values from the product_ids list

    def prompt_missing_word_docs(self, missing_docs):
        """
        Opens a window displaying a list of products for which Word documents are missing. 
        Provides options to create missing documents or exit the process.
        """

        # Log the start of the prompt for missing Word documents
        self.logger.info("Prompting for missing Word documents")

        self.correlate_window = Toplevel(self)
        self.correlate_window.title("Correlate Data")

        self.missing_docs = missing_docs

        # Create a Treeview with columns
        self.correlate_tree = ttk.Treeview(self.correlate_window, columns=('Folder Name', 'Product ID', 'Product Name'), show='headings')
        self.correlate_tree.pack(fill='both', expand=True)

        # Configure the columns
        self.correlate_tree.column('Folder Name', anchor='w', width=150)
        self.correlate_tree.column('Product ID', anchor='center', width=100)
        self.correlate_tree.column('Product Name', anchor='w', width=150)

        # Define the headings
        self.correlate_tree.heading('Folder Name', text='Folder Name', anchor='w')
        self.correlate_tree.heading('Product ID', text='Product ID', anchor='center')
        self.correlate_tree.heading('Product Name', text='Product Name', anchor='w')

        # Add the items to the Treeview
        for i, (folder_name, product_id, product_name) in enumerate(missing_docs):
            self.correlate_tree.insert('', 'end', iid=str(i), values=(folder_name, product_id, product_name))

        self.logger.info("Treeview setup completed with missing document entries")

        # Bind double-click event to an item
        self.correlate_tree.bind('<Double-1>', self.on_item_double_click)
        
        # Adding a Yes to All button
        yes_to_all_button = ttk.Button(self.correlate_window, text="Yes to All", command=self.create_all_word_docs)
        yes_to_all_button.pack()

        exit_button = ttk.Button(self.correlate_window, text="Exit", command=self.exit_correlate_window)
        exit_button.pack()

        self.logger.info("Correlation prompt window setup completed")

    def on_item_double_click(self, event):
        """
        Handles the double-click event on a tree item in the correlation window. 
        Initiates the creation of a Word document for the selected item.
        """

        # Log when an item is double-clicked
        self.logger.info("Item double-clicked in the correlation window")

        item_id = self.correlate_tree.selection()[0]
        item_values = self.correlate_tree.item(item_id, 'values')
        doc_data = (item_values[0], item_values[1], item_values[2])

        # Log the data of the item being processed
        self.logger.info(f"Processing item for Word document creation: {doc_data}")

        self.create_word_doc(doc_data, item_id)
    
    def exit_correlate_window(self):
        """
        Closes the window that shows Word documents that can be created for each product 
        and opens the Settings window.
        """
        self.logger.info("Closing the correlate window and opening the settings window")

        self.correlate_window.destroy()
        self.logger.info("Correlate window closed.")
        self.Settings_Window_Start()


    def products_to_sell_report(self):

        self.logger.info("Starting products to sell report generation")

        # Ensure the Excel file path and sheet name are set
        filepath, sheet_name = self.load_excel_path_and_sheet()
        if not filepath or not sheet_name:
            self.logger.error("Excel file path or sheet name is not set")
            messagebox.showerror("Error", "Excel file path or sheet name is not set.")
            return

        # Define the To Sell folder path
        to_sell_folder = self.to_sell_folder
        if not os.path.exists(to_sell_folder):
            self.logger.error("To Sell folder path is not set or does not exist")
            messagebox.showerror("Error", "To Sell folder path is not set or does not exist.")
            return

        # Check for existing folder starting with "- See products added on"
        folder_prefix = "- See products added on "
        existing_folder = None
        for folder in os.listdir(to_sell_folder):
            if folder.startswith(folder_prefix):
                existing_folder = folder
                break

        # Get today's date in the required format
        today_formatted = datetime.now().strftime("%Y-%m-%d")

        # Folder path for the new or existing folder
        new_folder_name = f"{folder_prefix}{today_formatted}"
        new_folder_path = os.path.join(to_sell_folder, new_folder_name)

        if existing_folder:
            # Rename the existing folder
            os.rename(os.path.join(to_sell_folder, existing_folder), new_folder_path)
        else:
            # Create a new folder
            os.makedirs(new_folder_path)

        self.logger.info("Loading data from the Excel workbook")

        # Load the original workbook and read the specified sheet into a DataFrame
        workbook = load_workbook(filepath, data_only=True)
        sheet = workbook[sheet_name]
        data = sheet.values
        columns = next(data)[0:]
        df = pd.DataFrame(data, columns=columns)

        # Get the names of the folders in the to_sell_folder and extract product IDs
        folder_names = os.listdir(to_sell_folder)
        folder_product_ids = set(folder_name.split(' ', 1)[0] for folder_name in folder_names)

        self.logger.info("Filtering and processing product data")

        # Filter out unwanted products and keep only necessary columns
        initial_count = len(df)
        df = df[(df['Damaged'] != 'YES') & (df['Cancelled Order'] != 'YES') & (df['Personal'] != 'YES') & (df['Sold'] != 'YES') & (~pd.isna(df['Product ID']))]
        df = df[df['Product ID'].isin(folder_product_ids)]
        df = df[['Product ID', 'To Sell After', 'Product Name', 'Product Price After IVU']]
        filtered_count = len(df)
        self.logger.info(f"Filtered from {initial_count} products to {filtered_count} products")

        # Convert 'To Sell After' to datetime
        df['To Sell After'] = pd.to_datetime(df['To Sell After'], errors='coerce')
        today = pd.to_datetime('today').normalize()
        df = df.dropna(subset=['To Sell After'])
        df = df[df['To Sell After'] <= today]
        self.logger.info("Converted 'To Sell After' dates and performed additional filtering")

        # Sort the DataFrame by 'Product ID'
        sorted_df = df.sort_values(by='Product ID', ascending=True)
        self.logger.info("Sorted the DataFrame based on 'Product ID'")

        # Call get_previous_excel_report_data and assign the return value to listx
        previous_product_ids, latest_file_date = self.get_previous_excel_report_data()
        self.logger.info(f"Retrieved data from the previous report dated {latest_file_date}")

        # Define the light green fill
        light_green_fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')

        # Create a new workbook and add the sorted data to it
        self.logger.info("Creating new workbook for the report")
        new_workbook = Workbook()
        new_sheet = new_workbook.active
        new_sheet.title = sheet_name

        self.logger.info("Applying formatting and styles to the workbook")
        for r_idx, row in enumerate(dataframe_to_rows(sorted_df, index=False, header=True), start=1):
            for c_idx, value in enumerate(row, start=1):
                cell = new_sheet.cell(row=r_idx, column=c_idx, value=value)
                if c_idx == 1 and r_idx > 1:  # Skip header row                    
                    if cell.value is not None and cell.value.upper() in previous_product_ids:
                        previous_product_ids.remove(cell.value)
                    else:
                        cell.fill = light_green_fill
                if c_idx == 2 and r_idx > 1:  # Skip header row
                    cell.number_format = 'MM/DD/YYYY'
                # Apply currency format to 'Product Price After IVU' column (assuming it's the fourth column)
                if c_idx == 4 and r_idx > 1:  # Skip header row
                    cell.number_format = '"$"#,##0.00'
                # Apply middle and center alignment to all cells
                if c_idx == 3 and r_idx > 1:  # 'Product Name' column
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                else:
                    cell.alignment = Alignment(horizontal='center', vertical='center')

        # Define the table dimensions
        table_ref = f"A1:{chr(65 + sorted_df.shape[1] - 1)}{sorted_df.shape[0] + 1}"

        # Create a table
        self.logger.info("Creating a table in the new workbook")
        table = Table(displayName="ProductsToSellTable", ref=table_ref)
        style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=True)
        table.tableStyleInfo = style
        new_sheet.add_table(table)

        # Adjust column widths
        new_sheet.column_dimensions['A'].width = 120 / 7  # Width for 'Product ID'
        new_sheet.column_dimensions['B'].width = 120 / 7  # Width for 'To Sell After'
        new_sheet.column_dimensions['C'].width = 700 / 7  # Width for 'Product Name'
        new_sheet.column_dimensions['D'].width = 200 / 7  # Width for 'Product Price After IVU'

        if latest_file_date is not None:
            formatted_date = latest_file_date.strftime('%A, %B %d, %Y')
            new_sheet['F2'] = f"Product IDs highlighted in green represent new products added since the \nlast report from {formatted_date}."
        else:
            new_sheet['F2'] = "Product IDs highlighted in green represent new products added."

        new_sheet['F3'] = datetime.now().strftime("This report was generated on %A, %B %d, %Y at %I:%M %p.")

        # Creating an Alignment object for center and middle alignment
        align_center_middle = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Applying the alignment and fill to the cells
        new_sheet['F2'].alignment = align_center_middle
        new_sheet['F2'].fill = light_green_fill
        new_sheet['F3'].alignment = align_center_middle
        new_sheet['F3'].fill = light_green_fill

        # Setting the width of column 'F' to 80 points
        new_sheet.column_dimensions['F'].width = 80


        self.logger.info("Finalizing and saving the report")


        self.logger.info("Saving the new workbook")
        today_str = datetime.now().strftime("%Y-%m-%d")
        new_report_path = os.path.join(new_folder_path, f"Products To Sell - {today_str}.xlsx")
        new_workbook.save(new_report_path)
        self.logger.info(f"Report saved at {new_report_path}")


        # Call the method to backup old reports
        self.backup_old_reports(new_folder_path, new_report_path)

        # Open the modified Excel file
        if sys.platform == "win32":
            os.startfile(new_report_path)
        elif sys.platform == "darwin":  # macOS
            subprocess.run(["open", copy_path])
        else:  # Linux variants
            subprocess.run(["xdg-open", copy_path])

    def get_previous_excel_report_data(self):
        self.logger.info("Starting to get previous Excel report data")

        to_sell_folder = self.to_sell_folder
        folder_prefix = "- See products added on "
        backup_folder_prefix = "Products to Sell Reports Backup"
        latest_file_date = None
        latest_file_path = None
        today = datetime.now().date()

        def find_latest_file(folder):
            nonlocal latest_file_date, latest_file_path
            for file in os.listdir(folder):
                if file.endswith(".xlsx") and file.startswith("Products To Sell -"):
                    file_date_str = file[len("Products To Sell - "):-len(".xlsx")]
                    try:
                        file_date = datetime.strptime(file_date_str, "%Y-%m-%d").date()
                        if file_date < today and (latest_file_date is None or file_date > latest_file_date):
                            latest_file_date = file_date
                            latest_file_path = os.path.join(folder, file)
                    except ValueError as e:
                        self.logger.error(f"Error parsing date from file name '{file}': {e}")

        # Check the latest file in the current folder
        current_folder_date = datetime.now().strftime("%Y-%m-%d")
        current_folder = os.path.join(to_sell_folder, f"{folder_prefix}{current_folder_date}")
        if os.path.exists(current_folder):
            find_latest_file(current_folder)

        # Check the backup folder if no file found in the current folder
        if latest_file_path is None:
            backup_folder = os.path.join(os.path.dirname(self.inventory_folder), "Excel Backups", backup_folder_prefix)
            if os.path.exists(backup_folder):
                find_latest_file(backup_folder)

        if latest_file_path is None:
            self.logger.info("No previous report found.")
            return [0], None

        self.logger.info(f"Previous report found at: {latest_file_path}")

        workbook = load_workbook(latest_file_path, data_only=True)
        sheet = workbook.active
        product_ids = [row[0] for row in sheet.iter_rows(min_row=2, values_only=True) if row[0] is not None]

        return set(product_ids), latest_file_date

    def backup_old_reports(self, current_report_folder, new_report_path):
        self.logger.info("Starting backup of old reports")

        # Define the backup folder path
        parent_dir = os.path.dirname(self.inventory_folder)
        backup_folder = os.path.join(parent_dir, "Excel Backups", "Products to Sell Reports Backup")

        # Create the backup folder if it doesn't exist
        if not os.path.exists(backup_folder):
            os.makedirs(backup_folder)
            self.logger.info(f"Backup folder '{backup_folder}' created.")
        else:
            self.logger.info(f"Backup folder '{backup_folder}' already exists.")

        # Maintain a maximum of 100 backups
        existing_backups = sorted(os.listdir(backup_folder))
        while len(existing_backups) >= 100:
            oldest_backup = existing_backups.pop(0)
            os.remove(os.path.join(backup_folder, oldest_backup))
            self.logger.info(f"Deleted oldest backup: {oldest_backup}")

        # Move the files
        for file in os.listdir(current_report_folder):
            file_path = os.path.join(current_report_folder, file)
            if file_path.endswith(".xlsx") and file_path != new_report_path:
                backup_path = os.path.join(backup_folder, file)
                shutil.move(file_path, backup_path)
                self.logger.info(f"Moved '{file}' to backup.")


# Product Form with functions used in it.
    def Product_Form(self):
        self.logger.info("Initializing product form widgets")
        try:

            # Create a style object
            style = ttk.Style()
            
            style.map('BlackOnDisabled.TEntry', foreground=[('disabled', 'black')])

            # Define a custom style named 'Blue.TButton' that changes the foreground color to blue
            style.configure('Blue.TButton', foreground='blue')

            # Create a custom font with a larger size
            link_font = Font(family="Helvetica", size=10)  # Adjust the size as per your requirement
            product_name_font = Font(family="Helvetica", size=11)  # Adjust the size as per your requirement

            # Add validation commands
            validate_percentage_command = (self.register(lambda P: self.validate_input(P, is_percentage=True)), '%P')
            validate_price_command = (self.register(self.validate_input), '%P')
            vcmd = (self.register(self.validate_input), '%P')


            self.product_frame = tk.Frame(self.bottom_frame, bg='light gray')
            self.product_frame.pack(side='right', fill='both', expand=True) #change pack to grid later

            # Row 0 Widgets
            self.row0_frame = tk.Frame(self.product_frame, bg='light gray')
            self.row0_frame.grid(row=0, column=5, sticky='ne', padx=50, pady=0)
        
            self.save_button = ttk.Button(self.row0_frame, text='Save', command=self.save, state='disabled')
            self.save_button.grid(row=0, column=0, sticky='w', padx=0, pady=0)

            self.edit_button = ttk.Button(self.row0_frame, text="Edit", command=self.toggle_edit_mode, state='disabled')
            self.edit_button.grid(row=0, column=1, sticky='w', padx=0, pady=0)


            # Row 1 Widgets
            self.row1_frame = tk.Frame(self.product_frame, bg='light gray')
            self.row1_frame.grid(row=1, column=0, sticky='nw', padx=5, pady=5)
            
            self.order_date_var = tk.StringVar()
            self.order_date_label = ttk.Label(self.row1_frame, text='Order Date')
            self.order_date_label.grid(row=0, column=0, sticky='w', padx=0, pady=0)
            self.order_date_entry = ttk.Entry(self.row1_frame, textvariable=self.order_date_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.order_date_entry.grid(row=1, column=0, sticky='w', padx=0, pady=0)

            self.to_sell_after_var = tk.StringVar()
            self.to_sell_after_label = ttk.Label(self.row1_frame, text='To Sell After')
            self.to_sell_after_label.grid(row=2, column=0, sticky='w', padx=0, pady=0)
            self.to_sell_after_entry = ttk.Entry(self.row1_frame, textvariable=self.to_sell_after_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.to_sell_after_entry.grid(row=3, column=0, sticky='w', padx=0, pady=0)

            # Row 1 Widgets
            # Column 1 Widget
            
            self.r1column1_frame = tk.Frame(self.product_frame, bg='light gray')
            self.r1column1_frame.grid(row=1, column=2, sticky='nw', padx=0, pady=0)
            self.product_image_label = ttk.Label(self.r1column1_frame, text='Image not loaded')
            self.product_image_label.grid(row=0, column=1, sticky='w', padx=0, pady=0)
            
            
            # Row 2 Widgets
            # Column 0 Widget
            # Create a new frame for the column 0 widgets
            self.r2column0_frame = tk.Frame(self.product_frame, bg='light gray')
            self.r2column0_frame.grid(row=2, column=0, sticky='nw', padx=25, pady=25)
            
            self.product_id_var = tk.StringVar()
            self.product_id_label = ttk.Label(self.r2column0_frame, text='Product ID')
            self.product_id_label.grid(row=0, column=0, sticky='w', padx=0, pady=0)
            self.product_id_entry = ttk.Entry(self.r2column0_frame, textvariable=self.product_id_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.product_id_entry.grid(row=1, column=0, sticky='w', padx=0, pady=0) 

            self.rack_id_var = tk.StringVar()
            self.rack_id_label = ttk.Label(self.r2column0_frame, text='Rack ID')
            self.rack_id_label.grid(row=0, column=0, sticky='e', padx=0, pady=0)
            self.rack_id_entry = ttk.Entry(self.r2column0_frame, textvariable=self.rack_id_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.rack_id_entry.grid(row=1, column=0, sticky='e', padx=0, pady=0)

            self.r2column0_frame.grid_rowconfigure(2, minsize=2)  # Adjust 'minsize' for desired space

            self.product_name_var = tk.StringVar()
            self.product_name_label = ttk.Label(self.r2column0_frame, text='Product Name')
            self.product_name_label.grid(row=3, column=0, sticky='w', padx=0, pady=0)

            # Create the Text widget with the desired background color inside the border frame
            self.product_name_text = tk.Text(self.r2column0_frame, height=8, width=50, bg="#eff0f1", fg="#000000", wrap="word", bd=0, highlightthickness=1, highlightcolor="#94cfeb", font=product_name_font)
            self.product_name_text.grid(row=4, column=0, sticky='w', padx=0, pady=1)
            
            # Bind the mouse click event to an empty lambda function
            self.product_name_text.bind("<Button-1>", lambda e: "break")
            
            self.r2column0_frame.grid_rowconfigure(5, minsize=2)  # Adjust 'minsize' for desired space
            
            self.product_folder_var = tk.StringVar()
            self.product_folder_label = ttk.Label(self.r2column0_frame, text='Product Folder')
            self.product_folder_label.grid(row=6, column=0, sticky='w', padx=0, pady=2)

            # Now use this style when creating your button
            self.product_folder_link = ttk.Button(self.r2column0_frame, textvariable=self.product_folder_var, style='Blue.TButton')

            self.product_folder_link.grid(row=7, column=0, sticky='w', padx=0, pady=0)

            self.r2column0_frame.grid_rowconfigure(8, minsize=2)  # Adjust 'minsize' for desired space

            self.order_link_var = tk.StringVar()
            self.order_link_label = ttk.Label(self.r2column0_frame, text='Order Link')
            self.order_link_label.grid(row=9, column=0, sticky='w', padx=0, pady=0)
            
            # Replace the Entry with a Text widget for clickable links
            self.order_link_text = tk.Text(self.r2column0_frame, height=1, width=40, bg="#eff0f1", fg="#000000", wrap=tk.NONE, bd=0, font=link_font)
            self.order_link_text.grid(row=10, column=0, sticky='w', padx=0, pady=1)
            self.order_link_text.tag_configure("hyperlink", foreground="blue", underline=True)
            self.order_link_text.bind("<Button-1>", self.open_hyperlink)
            self.order_link_text.config(state='disabled')

            self.r2column0_frame.grid_rowconfigure(11, minsize=2)  # Adjust 'minsize' for desired space

            self.asin_var = tk.StringVar()
            self.asin_label = ttk.Label(self.r2column0_frame, text='ASIN')
            self.asin_label.grid(row=12, column=0, sticky='w', padx=0, pady=0)
            self.asin_entry = ttk.Entry(self.r2column0_frame, textvariable=self.asin_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.asin_entry.grid(row=13, column=0, sticky='w', padx=0, pady=0)


            # Row 2 Widgets
            # Column 1 Widgets

            self.r2column1_frame = tk.Frame(self.product_frame, bg='light gray')
            self.r2column1_frame.grid(row=2, column=1, sticky='nw', padx=0, pady=5)
            custom_font = Font(family="Helvetica", size=7)
            style.configure('SmallFont.TButton', font=custom_font, padding=1)
            
            self.r2column1_frame.grid_rowconfigure(0, minsize=75)  # Adjust 'minsize' for desired space
            self.fair_market_value_var = tk.StringVar()
            self.fair_market_value_label = ttk.Label(self.r2column1_frame, text='Fair Market Value')
            self.fair_market_value_label.grid(row=2, column=0, sticky='w', padx=0, pady=0)
            self.fair_market_value_entry = ttk.Entry(self.r2column1_frame, textvariable=self.fair_market_value_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.fair_market_value_entry.grid(row=3, column=0, sticky='w', padx=0, pady=0)
            
            self.regular_product_price_var = tk.StringVar()
            self.regular_product_price_label = ttk.Label(self.r2column1_frame, text='Product Price')
            self.regular_product_price_label.grid(row=4, column=0, sticky='w', padx=0, pady=0)
            self.regular_product_price_entry = ttk.Entry(self.r2column1_frame, textvariable=self.regular_product_price_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.regular_product_price_entry.grid(row=5, column=0, sticky='w', padx=0, pady=0)
            
            self.ivu_tax_var = tk.StringVar()
            self.ivu_tax_label = ttk.Label(self.r2column1_frame, text='IVU Tax')
            self.ivu_tax_label.grid(row=6, column=0, sticky='w', padx=0, pady=0)
            self.ivu_tax_entry = ttk.Entry(self.r2column1_frame, textvariable=self.ivu_tax_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.ivu_tax_entry.grid(row=7, column=0, sticky='w', padx=0, pady=0)
            
            self.product_price_plus_ivu_var = tk.StringVar()
            self.product_price_plus_ivu_label = ttk.Label(self.r2column1_frame, text='Product Price (+ IVU)')
            self.product_price_plus_ivu_label.grid(row=8, column=0, sticky='w', padx=0, pady=0)
            self.product_price_plus_ivu_entry = ttk.Entry(self.r2column1_frame, textvariable=self.product_price_plus_ivu_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.product_price_plus_ivu_entry.grid(row=9, column=0, sticky='w', padx=0, pady=0)

            # Row 2 Widgets
            # Column 2 Widgets

            self.r2column2_frame = tk.Frame(self.product_frame, bg='light gray')
            self.r2column2_frame.grid(row=2, column=2, sticky='nw', padx=0, pady=5)
            custom_font = Font(family="Helvetica", size=7)
            style.configure('SmallFont.TButton', font=custom_font, padding=1)
            
            self.r2column2_frame.grid_rowconfigure(0, minsize=75)  # Adjust 'minsize' for desired space

            self.discount_var = tk.StringVar()
            self.discount_label = ttk.Label(self.r2column2_frame, text='Discount($ Or %)')
            self.discount_label.grid(row=1, column=0, sticky='w', padx=0, pady=0)

            # Frame to hold the discount entries
            self.discount_frame = ttk.Frame(self.r2column2_frame)
            self.discount_frame.grid(row=2, column=0, sticky='w', padx=0, pady=0)

            # Discount entries with validation and event binding
            self.discount_var = tk.StringVar()
            self.discount_entry = ttk.Entry(self.discount_frame, textvariable=self.discount_var, width=8, state='disabled', style='BlackOnDisabled.TEntry', validate='key', validatecommand=validate_price_command)
            self.discount_entry.pack(side=tk.LEFT)
            #self.discount_entry.bind("<KeyRelease>", self.on_price_changed)        
            self.discount_entry.bind("<FocusIn>", self.on_discount_price_focus_in)        
            self.discount_entry.bind("<FocusOut>", self.on_discount_price_focus_out)

            # Label "Or"
            self.or_label = ttk.Label(self.discount_frame, text="Or")
            self.or_label.pack(side=tk.LEFT)

            self.percent_discount_var = tk.StringVar()
            self.percent_discount_entry = ttk.Entry(self.discount_frame, textvariable=self.percent_discount_var, width=8, state='disabled', style='BlackOnDisabled.TEntry', validate='key', validatecommand=validate_percentage_command)
            self.percent_discount_entry.pack(side=tk.LEFT)
            #self.percent_discount_entry.bind("<KeyRelease>", self.on_percentage_changed)
            self.percent_discount_entry.bind("<FocusIn>", self.on_discount_percentage_focus_in)
            self.percent_discount_entry.bind("<FocusOut>", self.on_discount_percentage_focus_out)
            
            self.product_price_after_discount_var = tk.StringVar()
            self.product_price_after_discount_label = ttk.Label(self.r2column2_frame, text='Product Price after Discount')
            self.product_price_after_discount_label.grid(row=3, column=0, sticky='w', padx=0, pady=0)
            self.product_price_after_discount_entry = ttk.Entry(self.r2column2_frame, textvariable=self.product_price_after_discount_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.product_price_after_discount_entry.grid(row=4, column=0, sticky='w', padx=0, pady=0)

            self.ivu_tax_after_discount_var = tk.StringVar()
            self.ivu_tax_after_discount_label = ttk.Label(self.r2column2_frame, text='IVU Tax after Discount')
            self.ivu_tax_after_discount_label.grid(row=5, column=0, sticky='w', padx=0, pady=0)
            self.ivu_tax_after_discount_entry = ttk.Entry(self.r2column2_frame, textvariable=self.ivu_tax_after_discount_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.ivu_tax_after_discount_entry.grid(row=6, column=0, sticky='w', padx=0, pady=0)

            self.product_price_minus_discount_plus_ivu_var = tk.StringVar()
            self.product_price_minus_discount_plus_ivu_label = ttk.Label(self.r2column2_frame, text='Product Price (+IVU - Discount)')
            self.product_price_minus_discount_plus_ivu_label.grid(row=7, column=0, sticky='w', padx=0, pady=0)
            self.product_price_minus_discount_plus_ivu_entry = ttk.Entry(self.r2column2_frame, textvariable=self.product_price_minus_discount_plus_ivu_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.product_price_minus_discount_plus_ivu_entry.grid(row=8, column=0, sticky='w', padx=0, pady=0)

            self.sold_date_var = tk.StringVar()
            self.sold_date_label = ttk.Label(self.r2column2_frame, text='Sold Date')
            self.sold_date_label.grid(row=9, column=0, sticky='w', padx=0, pady=0)
            
            self.sold_date_entry = ttk.Entry(self.r2column2_frame, textvariable=self.sold_date_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.sold_date_entry.grid(row=10, column=0, sticky='w', padx=0, pady=0)
            
            self.sold_date_button = ttk.Button(self.r2column2_frame, text="Pick\nDate", style='SmallFont.TButton', command=self.pick_date, state='disabled', width=5)
            self.sold_date_button.grid(row=10, column=0, sticky='e', padx=0, pady=0)

            self.clear_button = ttk.Button(self.r2column2_frame, text="Clear\nDate", style='SmallFont.TButton', command=self.clear_date, state='disabled', width=5)
            self.clear_button.grid(row=10, column=1, sticky='e', padx=0, pady=0)

            self.payment_type_var = tk.StringVar()
            self.payment_type_label = ttk.Label(self.r2column2_frame, text='Payment Type')
            self.payment_type_label.grid(row=11, column=0, sticky='w', padx=0, pady=0)
            
            self.payment_type_combobox = ttk.Combobox(self.r2column2_frame, textvariable=self.payment_type_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.payment_type_combobox['values'] = ('', 'Cash', 'ATH Movil')
            self.payment_type_combobox.grid(row=12, column=0, sticky='w', padx=0, pady=0)        
            
            self.sold_price_var = tk.StringVar()
            self.sold_price_label = ttk.Label(self.r2column2_frame, text='Sold Price')
            self.sold_price_label.grid(row=13, column=0, sticky='w', padx=0, pady=0)
            self.sold_price_entry = ttk.Entry(self.r2column2_frame, textvariable=self.sold_price_var, state='disabled', style='BlackOnDisabled.TEntry')
            self.sold_price_entry.grid(row=14, column=0, sticky='w', padx=0, pady=0)


            # Row 2 Widgets
            # Column 3 Widgets
            # Creating a new frame for checkboxes within the product frame
            self.checkbox_frame = tk.Frame(self.product_frame, bg='light gray')
            self.checkbox_frame.grid(row=2, column=3, rowspan=8, sticky='nw', padx=0, pady=5)
            self.checkbox_frame.grid_rowconfigure(0, minsize=75)  # Adjust 'minsize' for desired space

            self.sold_var = tk.BooleanVar()
            self.sold_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Sold', variable=self.sold_var)
            self.sold_checkbutton.grid(row=1, column=0, sticky='w', padx=0, pady=0)
            
            self.checkbox_frame.grid_rowconfigure(2, minsize=20)  # This creates a 20-pixel-high empty row as a spacer
            
            self.cancelled_order_var = tk.BooleanVar()
            self.cancelled_order_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Cancelled Order', variable=self.cancelled_order_var)
            self.cancelled_order_checkbutton.grid(row=3, column=0, sticky='w', padx=0, pady=0)

            self.damaged_var = tk.BooleanVar()
            self.damaged_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Damaged', variable=self.damaged_var)
            self.damaged_checkbutton.grid(row=4, column=0, sticky='w', padx=0, pady=0)

            self.personal_var = tk.BooleanVar()
            self.personal_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Personal', variable=self.personal_var)
            self.personal_checkbutton.grid(row=5, column=0, sticky='w', padx=0, pady=0)

            self.checkbox_frame.grid_rowconfigure(6, minsize=20)  # This creates a 20-pixel-high empty row as a spacer

            self.reviewed_var = tk.BooleanVar()
            self.reviewed_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Reviewed', variable=self.reviewed_var)
            self.reviewed_checkbutton.grid(row=7, column=0, sticky='w', padx=0, pady=0)

            self.pictures_downloaded_var = tk.BooleanVar()
            self.pictures_downloaded_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Pictures Downloaded', variable=self.pictures_downloaded_var)
            self.pictures_downloaded_checkbutton.grid(row=8, column=0, sticky='w', padx=0, pady=0)

            self.uploaded_to_site_var = tk.BooleanVar()
            self.uploaded_to_site_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Uploaded to Site', variable=self.uploaded_to_site_var)
            self.uploaded_to_site_checkbutton.grid(row=9, column=0, sticky='w', padx=0, pady=0)

            self.product_frame.grid_rowconfigure(3, minsize=60)  # This creates a 20-pixel-high empty row as a spacer
            

            # Row 4 Widgets
            # Column 0 Widgets
            # Creating a new frame for checkboxes within the product frame
            self.comments_frame = tk.Frame(self.product_frame, bg='light gray')
            self.comments_frame.grid(row=4, column=0, columnspan=3, sticky='nw', padx=25, pady=5)

            self.comments_text = tk.Text(self.comments_frame, height=8, width=150, bg="#eff0f1", fg="#000000", wrap="word", state="disabled", bd=0, highlightthickness=1, highlightcolor="#94cfeb", font=product_name_font)
            self.comments_text.grid(row=0, column=0, sticky='w', padx=0, pady=1)

            self.comments_text.bind("<FocusIn>", self.on_comments_focus_in)
            self.comments_text.bind("<FocusOut>", self.on_comments_focus_out)

            self.product_description_text = tk.Text(self.comments_frame, height=8, width=150, bg="#eff0f1", fg="#000000", wrap="word", state="disabled", bd=0, highlightthickness=1, highlightcolor="#94cfeb", font=product_name_font)
            self.product_description_text.grid(row=1, column=0, sticky='w', padx=0, pady=1)

            self.product_description_text.bind("<FocusIn>", self.on_product_description_focus_in)
            self.product_description_text.bind("<FocusOut>", self.on_product_description_focus_out)


            # Bind the new checkbox click control function to the checkboxes
            self.sold_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.sold_var))
            self.cancelled_order_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.cancelled_order_var))
            self.damaged_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.damaged_var))
            self.personal_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.personal_var))
            self.reviewed_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.reviewed_var))
            self.pictures_downloaded_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.pictures_downloaded_var))
            self.uploaded_to_site_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.uploaded_to_site_var))

            self.product_price_plus_ivu_entry.bind("<FocusIn>", self.on_price_focus_in)
            self.product_price_plus_ivu_entry.bind("<FocusOut>", self.on_price_focus_out)

            self.sold_price_entry.bind("<FocusIn>", self.on_price_focus_in)
            self.sold_price_entry.bind("<FocusOut>", self.on_price_focus_out)

            # configure validation commands
            self.fair_market_value_entry.config(validate='key', validatecommand=vcmd)
            self.regular_product_price_entry.config(validate='key', validatecommand=vcmd)
            self.ivu_tax_entry.config(validate='key', validatecommand=validate_price_command)
            self.product_price_plus_ivu_entry.config(validate='key', validatecommand=vcmd)
            self.product_price_after_discount_entry.config(validate='key', validatecommand=vcmd)
            self.ivu_tax_after_discount_entry.config(validate='key', validatecommand=vcmd)
            self.product_price_minus_discount_plus_ivu_entry.config(validate='key', validatecommand=vcmd)
            self.sold_price_entry.config(validate='key', validatecommand=vcmd)

            # Load settings
            try:
                with open("folders_paths.txt", "r") as file:
                    lines = file.read().splitlines()
                    self.inventory_folder = lines[0]
                    self.sold_folder = lines[1]
                    self.to_sell_folder = lines[2] if len(lines) > 2 else None
                    if self.inventory_folder:  # Check if inventory_folder is defined
                        self.combine_and_display_folders()
            except FileNotFoundError:
                pass
            self.logger.info("Product form initialized")

        except Exception as e:
            self.logger.error(f"Error initializing main window widgets: {e}")

    def display_product_details(self, event):
        """
        Displays the details of a selected product in the GUI. The details are fetched 
        from an Excel sheet based on the selected product ID. This function updates various 
        fields in the GUI with the product's information.
        """

        # Before fetching product details
        self.logger.info("Displaying product details")

        selection = self.folder_list.curselection()
        # Get the index of the selected item
        if not selection:
            return  # No item selected
        index = selection[0]
        selected_folder_name = self.folder_list.get(index)
        selected_product_id = selected_folder_name.split(' ')[0].upper()  # Assuming the product ID is at the beginning
        self.current_product_id = selected_product_id  # Set the current product ID

        if self.edit_mode:
            self.toggle_edit_mode()

        # Ensure that the Excel file path and sheet name are set
        filepath, sheet_name = self.load_excel_path_and_sheet()
        if filepath and sheet_name:
            self.excel_manager.filepath = filepath
            self.excel_manager.sheet_name = sheet_name
            self.excel_manager.load_data()  # Load the data

            # Retrieve product information from the DataFrame
            try:
                product_info = self.excel_manager.get_product_info(selected_product_id)
                # Right after fetching product_info
                self.product_folder_path = self.get_folder_path_from_db(selected_product_id)

                if product_info:

                    self.edit_button.config(state="normal")
                    self.order_link_text.config(state='normal')
                    self.cancelled_order_var.set(self.excel_value_to_bool(product_info.get('Cancelled Order')))
                    self.damaged_var.set(self.excel_value_to_bool(product_info.get('Damaged')))
                    self.personal_var.set(self.excel_value_to_bool(product_info.get('Personal')))
                    self.reviewed_var.set(self.excel_value_to_bool(product_info.get('Reviewed')))
                    self.pictures_downloaded_var.set(self.excel_value_to_bool(product_info.get('Pictures Downloaded')))
                    self.uploaded_to_site_var.set(self.excel_value_to_bool(product_info.get('Uploaded to Site')))
                    self.sold_var.set(self.excel_value_to_bool(product_info.get('Sold')))
                    
                    # For each field, check if the value is NaN using pd.isnull and set it to an empty string if it is
                    self.asin_var.set('' if pd.isnull(product_info.get('ASIN')) else product_info.get('ASIN', ''))
                    self.product_id_var.set('' if pd.isnull(product_info.get('Product ID')) else product_info.get('Product ID', ''))
                    self.rack_id_var.set(product_info.get('Rack ID', ''))


                    self.product_name_text.configure(state='normal')
                    self.product_name_text.delete(1.0, "end")
                    product_name = product_info.get('Product Name', '')
                    if product_name:
                        self.product_name_text.insert("insert", product_name) 
                    self.product_name_text.configure(state='disabled')

                    self.comments_text.configure(state='normal')
                    self.comments_text.delete(1.0, "end")

                    # Assuming product_info['Comments'] can be NaN, None, or a string
                    comments_text = product_info.get('Comments', None)
                    # Check for NaN (using pandas' isna function if you're working with pandas)
                    # You can also directly check if comments_text is None, which covers both None and NaN cases
                    if comments_text is None or pd.isna(comments_text):
                        display_text = "No Comments Found."
                    else:
                        display_text = comments_text
                    self.comments_text.insert("insert", display_text)
                    self.comments_text.configure(state='disabled')


                    self.product_description_text.configure(state='normal')
                    self.product_description_text.delete(1.0, "end")

                    product_description_text = product_info.get('Product Description', None)
                    # Check for NaN (using pandas' isna function if you're working with pandas)
                    # You can also directly check if product_description_text is None, which covers both None and NaN cases
                    if product_description_text is None or pd.isna(product_description_text):
                        display_product_description_text = "No Product Description At The Moment."
                    else:
                        display_product_description_text = product_description_text
                    self.product_description_text.insert("insert", display_product_description_text)
                    self.product_description_text.configure(state='disabled')

                    # When a product is selected and the order date is fetched
                    order_date = product_info.get('Order Date', '')
                    formatted_order_date = ''  # Default value
                    if isinstance(order_date, datetime):
                        formatted_order_date = order_date.strftime('%m/%d/%Y')
                        self.order_date_var.set(formatted_order_date)
                    elif isinstance(order_date, str) and order_date:
                        try:
                            # If the date is in the format 'mm/dd/yy', such as '2/15/23'
                            order_date = datetime.strptime(order_date, "%m/%d/%Y")
                            formatted_order_date = order_date.strftime('%m/%d/%Y')
                            self.order_date_var.set(formatted_order_date)
                        except ValueError as e:
                            
                            messagebox.showerror("Error", f"Incorrect date format: {e}")
                    else:
                        self.order_date_var.set('')
                    self.order_date_var.set(formatted_order_date)
                    
                    # When a product is selected and the order date is fetched
                    to_sell_after = product_info.get('To Sell After', '')
                    formatted_to_sell_after = ''  # Default value
                    if pd.notnull(to_sell_after):  # Check if 'To Sell After' is not null
                        try:
                            if isinstance(to_sell_after, datetime):
                                formatted_to_sell_after = to_sell_after.strftime('%m/%d/%Y')
                            elif isinstance(to_sell_after, str) and to_sell_after:
                                to_sell_after = datetime.strptime(to_sell_after, "%m/%d/%Y")
                                formatted_to_sell_after = to_sell_after.strftime('%m/%d/%Y')
                        except ValueError as e:
                            messagebox.showerror("Error", f"Incorrect date format: {e}")
                    self.to_sell_after_var.set(formatted_to_sell_after)
                    self.update_to_sell_after_color()
                    
                    sold_date = product_info.get('Sold Date', '')
                    formatted_sold_date = ''  # Default value

                    if pd.notnull(sold_date):  # Check if 'Sold Date' is not null
                        try:
                            if isinstance(sold_date, datetime):
                                formatted_sold_date = sold_date.strftime('%m/%d/%Y')
                            elif isinstance(sold_date, str) and sold_date:
                                # Parse the date string to a date object and format it
                                sold_date = datetime.strptime(sold_date, "%m/%d/%Y").date()
                                formatted_sold_date = sold_date.strftime('%m/%d/%Y')
                        except ValueError as e:
                            messagebox.showerror("Error", f"Incorrect date format: {e}")
                    else:
                        formatted_sold_date = ''

                    self.sold_date_var.set(formatted_sold_date)

                    def format_price(value):
                        if pd.isnull(value):
                            return ''
                        # Separate the fractional and integer parts
                        fractional, integer = math.modf(value)
                        # If the fractional part is 0, use the integer part; otherwise, format with two decimal places
                        return f"${int(integer) if fractional == 0 else f'{value:.2f}'}"

                    def format_percentage(value):
                        if pd.isnull(value):
                            return ''
                        # Separate the fractional and integer parts
                        fractional, integer = math.modf(value)
                        # If the fractional part is 0, use the integer part; otherwise, format with two decimal places
                        return f"{int(integer) if fractional == 0 else f'{value:.2f}'}%"

                    self.fair_market_value_var.set(format_price(product_info.get('Fair Market Value')))
                    self.discount_var.set(format_price(product_info.get('Discount')))
                    self.percent_discount_var.set(format_percentage(product_info.get('Discount Percentage')))

                    self.regular_product_price_var.set(format_price(product_info.get('Product Price')))
                    self.ivu_tax_var.set(format_price(product_info.get('IVU Tax')))
                    self.product_price_plus_ivu_var.set(format_price(product_info.get('Product Price After IVU')))

                    self.product_price_after_discount_var.set(format_price(product_info.get('Product Price After Discount')))
                    self.ivu_tax_after_discount_var.set(format_price(product_info.get('IVU Tax After Discount')))
                    self.product_price_minus_discount_plus_ivu_var.set(format_price(product_info.get('Product Price After IVU and Discount')))

                    self.sold_price_var.set(format_price(product_info.get('Sold Price')) if not pd.isnull(product_info.get('Sold Price')) else '')

                    self.order_link_text.delete(1.0, "end")
                    hyperlink = product_info.get('Order Link', '')
                    if hyperlink:
                        self.order_link_text.insert("insert", hyperlink, "hyperlink")
                        self.order_link_text.tag_add("hyperlink", "1.0", "end")
                        
                    self.payment_type_var.set('' if pd.isnull(product_info.get('Payment Type')) else product_info.get('Payment Type', ''))
                    # ... continue with other fields as needed ...
                    # Add code here to populate the Sold Date and other date-related fields, if applicable
                    
                    # Fetch the full folder path from the database using the product ID.
                    folder_path = self.get_folder_path_from_db(selected_product_id)

                    # Extract the name of the parent directory (where the product folder is located)
                    parent_folder_name = os.path.basename(os.path.dirname(folder_path)) if folder_path else "No Folder"
                    self.product_folder_var.set(parent_folder_name)

                    # If the folder path exists, update the button to open the product folder when clicked
                    if folder_path and os.path.exists(folder_path):
                        self.product_folder_link.config(command=lambda: self.open_product_folder(folder_path), state='normal')
                    else:
                        self.product_folder_var.set("No Folder")
                        self.product_folder_link.config(state='disabled')
                    self.product_image_label.config(image='')
                    self.product_image_label.configure(text='Loading image...')

                    # 1. Find the column number for "Product Image"
                    product_image_col_num = None
                    for col_num, col_name in enumerate(self.excel_manager.data_frame.columns):
                        if col_name == 'Product Image':
                            product_image_col_num = col_num
                            break
                    # 2. Get the current row number
                    current_row_num = self.excel_manager.data_frame[self.excel_manager.data_frame['Product ID'].str.upper() == selected_product_id.upper()].index[0]
  
                    # 3. Print the column name and row number
                    if product_image_col_num is not None:
                        self.load_and_display_image(current_row_num + 1, product_image_col_num, selected_product_id)
                    
                    self.logger.info(f"Product details displayed for: {selected_product_id}")
                else:
                    self.edit_button.config(state='disabled')
                    self.cancelled_order_var.set(False)
                    self.damaged_var.set(False)
                    self.personal_var.set(False)
                    self.reviewed_var.set(False)
                    self.pictures_downloaded_var.set(False)
                    self.uploaded_to_site_var.set(False)
                    
                    self.sold_var.set(False)
                    self.product_image_label.config(image='')
                    self.product_image_label.configure(text="Image not loaded.")
                    # Populate the widgets with the matched data
                    self.asin_var.set('')
                    self.product_id_var.set('')
                    self.rack_id_var.set('')
                    self.to_sell_after_var.set('')
                    # Add code here to handle the product image, if applicable
                    self.product_name_text.configure(state='normal')
                    self.product_name_text.delete(1.0, tk.END)
                    self.product_name_text.insert(tk.END, 'Product not found in Excel.')
                    self.product_name_text.configure(state='disabled')
                    self.comments_text.configure(state='normal')
                    self.comments_text.delete(1.0, tk.END)
                    self.comments_text.insert(tk.END, 'Comment not found in Excel.')
                    self.comments_text.configure(state='disabled')
                    
                    self.product_description_text.configure(state='normal')
                    self.product_description_text.delete(1.0, tk.END)
                    self.product_description_text.insert(tk.END, 'Product description not found in Excel.')
                    self.product_description_text.configure(state='disabled')
                    self.order_date_var.set('')
                    self.fair_market_value_var.set('')
                    self.discount_var.set('')
                    self.percent_discount_var.set("")

                    self.product_price_after_discount_var.set("")
                    self.ivu_tax_after_discount_var.set("")
                    self.product_price_minus_discount_plus_ivu_var.set("")

                    self.product_price_plus_ivu_var.set('')
                    self.ivu_tax_var.set('')
                    self.regular_product_price_var.set('')
                    self.order_link_text.delete(1.0, "end")
                    self.sold_price_var.set('')
                    self.payment_type_var.set('')
                    self.sold_date_var.set('')

                    # Fetch the full folder path from the database using the product ID.
                    folder_path = self.get_folder_path_from_db(selected_product_id)

                    # Extract the name of the parent directory (where the product folder is located)
                    parent_folder_name = os.path.basename(os.path.dirname(folder_path)) if folder_path else "No Folder"
                    self.product_folder_var.set(parent_folder_name)

                    # If the folder path exists, update the button to open the product folder when clicked
                    if folder_path and os.path.exists(folder_path):
                        self.product_folder_link.config(command=lambda: self.open_product_folder(folder_path), state='normal')
                    else:
                        self.product_folder_var.set("No Folder")
                        self.product_folder_link.config(state='disabled')

                    self.order_link_text.config(state='disabled')

            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {e}")
                #print(f"Error retrieving product details: {e}")
                self.logger.error(f"Error retrieving product details: {e}")
        else:
            messagebox.showerror("Error", "Excel file path or sheet name is not set.")        
        
        # Unbind the Enter key from the save_button's command
        self.master.unbind('<Return>')
        self.master.unbind('<Escape>')
       
        # Bind the Enter key to the global enter handler
        self.master.bind('<Return>', self.edit_on_key_handler)
        self.logger.info("Completed displaying product details")

    def refresh_and_select_product(self, product_id):
        """
        Refreshes the list of products and selects the specified product. 
        This function is typically called after updating product information to reflect changes in the UI.
        """

        # Log before starting the refresh process
        self.logger.info(f"Refreshing and selecting product with ID: {product_id}")

        # Refresh the list of products
        self.combine_and_display_folders()
        
        # Convert the product_id to uppercase for case-insensitive comparison
        product_id_upper = product_id.upper()

        # Log after refreshing the list
        self.logger.info("Product list refreshed")

        # Find the index of the product that was just edited
        product_index = None
        for index, product_name in enumerate(self.folder_list.get(0, tk.END)):
            # Use .split() to get the first part of the folder name and compare it in uppercase
            if product_name.split()[0].upper() == product_id_upper:
                product_index = index
                break
        
        # If the product is found in the list, select it
        if product_index is not None:
            self.folder_list.selection_set(product_index)
            self.folder_list.see(product_index)  # Ensure the product is visible in the list
            self.folder_list.event_generate("<<ListboxSelect>>")  # Trigger the event to display product details
            self.logger.info(f"Selected and displayed details for product ID: {product_id}")
            
        self.logger.info("Completed product selection process")
        self.toggle_edit_mode()

    def on_product_description_focus_in(self, event):
        # Assuming self.product_description_text is a text widget, 
        # you should get its current text using the appropriate method.
        current_text = self.product_description_text.get("1.0", "end-1c")
        if current_text == "No Product Description At The Moment.":
            self.product_description_text.delete("1.0", "end")

    def on_product_description_focus_out(self, event):
        # Get the current text from the widget.
        current_text = self.product_description_text.get("1.0", "end-1c")
        if current_text.strip() == "":
            self.product_description_text.insert("1.0", "No Product Description At The Moment.")

    def on_comments_focus_in(self, event):
        # Assuming self.product_description_text is a text widget, 
        # you should get its current text using the appropriate method.
        current_text = self.comments_text.get("1.0", "end-1c")
        if current_text == "No Comments Found.":
            self.comments_text.delete("1.0", "end")

    def on_comments_focus_out(self, event):
        # Get the current text from the widget.
        current_text = self.comments_text.get("1.0", "end-1c")
        if current_text.strip() == "":
            self.comments_text.insert("1.0", "No Comments Found.")

    def validate_input(self, input_value, is_percentage=False):
        # Check for empty input
        if input_value == "":
            return True

        # Check for more than one decimal point
        if input_value.count('.') > 1:
            return False

        # Split input on decimal point and check for two decimal places
        parts = input_value.split('.')
        if len(parts) == 2 and len(parts[1]) > 2:
            return False

        # Check if all characters are digits or a decimal point
        if all(ch.isdigit() or ch == '.' for ch in input_value):
            self.logger.info("Input is valid")
            return True
        else:
            self.logger.error("Input contains invalid characters")
            return False

    def on_price_focus_in(self, event):
        """Stores the initial price when focus is gained."""
        self.logger.info("Handling focus in event for price entry")
        entry_widget = event.widget
        price_str = entry_widget.get()
        try:
            if event.widget == self.product_price_plus_ivu_entry:
                self.initial_product_price_plus_ivu = price_str.lstrip('$')
            if price_str.startswith('$'):
                entry_widget.delete(0, tk.END)
                entry_widget.insert(0, price_str.lstrip('$'))
        except Exception as e:
            self.logger.error(f"Error handling on price focus in: {e}")

    def on_price_focus_out(self, event):
        if self.edit_mode:
            if self.trigger_price_focus_out_flag:
                entry_widget = event.widget
                current_price = entry_widget.get().lstrip('$')

                entry_widget.config(validate='none')

                try:
                    # Check if the widget is sold_price_entry
                    if entry_widget == self.sold_price_entry:
                        if current_price.strip():
                            try:
                                # Check if it's a valid number
                                float(current_price)
                                # If it's a number, add the $ sign
                                entry_widget.delete(0, tk.END)
                                entry_widget.insert(0, f"${current_price}")
                            except ValueError:
                                # If not a valid number, do not modify the value
                                pass
                        # For sold_price_entry, no further action is needed
                        return
                    # Convert current_price to a float for comparison
                    current_price_float = float(current_price) if current_price.strip() else 0.0

                    # Check if the price string is empty and set it and related fields to default values
                    if not current_price.strip() or current_price_float == 0.0:
                        self.logger.info("Price string is empty, resetting to default values")
                        entry_list = [self.regular_product_price_entry, self.ivu_tax_entry, 
                                    self.discount_entry, self.product_price_after_discount_entry, 
                                    self.ivu_tax_after_discount_entry, 
                                    self.product_price_minus_discount_plus_ivu_entry, 
                                    self.percent_discount_entry]

                        for widget in entry_list:
                            widget.config(validate='none', state='normal')

                            # Set the current widget to $0
                            self.product_price_plus_ivu_entry.delete(0, tk.END)
                            self.product_price_plus_ivu_entry.insert(0, "$0")

                            # Set related fields to their default values

                            self.regular_product_price_entry.delete(0, tk.END)
                            self.regular_product_price_entry.insert(0, "$0")

                            self.ivu_tax_entry.delete(0, tk.END)
                            self.ivu_tax_entry.insert(0, "$0")

                            # Temporarily disable validation and change state to normal
                            self.discount_entry.config(validate='none', state='normal')

                            # Set the widget to '$0'
                            self.discount_entry.delete(0, tk.END)
                            self.discount_entry.insert(0, "$0")

                            # Re-enable validation and change state back to disabled
                            self.discount_entry.config(validate='key', state='disabled')


                            self.product_price_after_discount_entry.delete(0, tk.END)
                            self.product_price_after_discount_entry.insert(0, "$0")

                            self.ivu_tax_after_discount_entry.delete(0, tk.END)
                            self.ivu_tax_after_discount_entry.insert(0, "$0")

                            self.product_price_minus_discount_plus_ivu_entry.delete(0, tk.END)
                            self.product_price_minus_discount_plus_ivu_entry.insert(0, "$0")

                            self.percent_discount_entry.delete(0, tk.END)
                            self.percent_discount_entry.insert(0, "0%")

                            # Re-enable validation and update widget state, disable all except specific widgets
                            for widget in entry_list:
                                if widget not in [self.discount_entry, self.percent_discount_entry, self.product_price_plus_ivu_entry]:
                                    widget.config(validate='key', state='disabled')
                                else:
                                    widget.config(validate='key', state='normal')  # Keep these widgets editable
                        # Update GUI
                        self.update_idletasks()
                        self.logger.info("Reset all price related fields to default values")
                        return 

                    if not current_price.startswith('$'):
                        entry_widget.delete(0, tk.END)
                        entry_widget.insert(0, f"${current_price}")
                        self.logger.info("Added dollar sign to the price entry")

                    entry_widget.config(validate='key')

                    # Handling change in product price plus IVU
                    if event.widget == self.product_price_plus_ivu_entry and self.initial_product_price_plus_ivu != current_price:
                        self.logger.info("Product price plus IVU changed, recalculating discounts")
                        if not hasattr(self, 'prompt_shown'):
                            self.prompt_shown = True

                            self.recalculate_original_price_and_tax()
                            discount_price = self.discount_var.get().lstrip('$')
                            discount_percentage = self.percent_discount_var.get().rstrip('%')
                            message = f"Product price changed. Calculate discount based on?\n\nPrice: ${discount_price}\nPercentage: {discount_percentage}%"
                            response = messagebox.askquestion("Discount Calculation", message)

                            if response == 'yes':
                                self.calculate_discount('price')
                            else:
                                self.calculate_discount('percentage')

                            del self.prompt_shown

                        self.initial_product_price_plus_ivu = ''

                except Exception as e:
                    self.logger.error(f"Error during focus out event processing: {e}")

    def save_on_key_handler(self, event):
        self.logger.info("Handling key press event for saving")

        try:
            if event.widget in [self.discount_entry, self.percent_discount_entry, self.product_price_plus_ivu_entry]:
                # Flag set for triggering save operation
                # self.trigger_save_flag = True 
                self.product_id_entry.focus_set()
                self.logger.info("Focus set to product ID entry, preparing for save operation")
            else:
                self.save()
                self.logger.info("Save function called directly due to key press on a different widget")
        except Exception as e:
            self.logger.error(f"Error in save on key handler: {e}")

    def edit_on_key_handler(self, event):
        self.logger.info("Handling key press event for edit mode")
        try:
            productid = self.product_id_entry.get().upper()

            if event.widget in [self.discount_entry, self.percent_discount_entry, self.product_price_plus_ivu_entry]:
                self.trigger_price_focus_out_flag = False
                self.search_entry.focus_set()
                self.trigger_price_focus_out_flag = True
                self.refresh_and_select_product(productid)
                self.logger.info(f"Edit mode toggled for product ID: {productid}")
                self.toggle_edit_mode()
            else:
                self.toggle_edit_mode()
                self.logger.info("Edit mode toggled for a different widget")
        except Exception as e:
            self.logger.error(f"Error in edit on key handler: {e}")

    def on_price_changed(self, event=None):
        self.last_changed = 'price'
        self.calculate_discount()

    def on_discount_price_focus_in(self, event=None):
        self.logger.info("Handling focus in event for discount price entry")

        try:
            price_str = self.discount_var.get()
            if price_str.startswith('$'):
                price_str = price_str.lstrip('$')
                self.discount_var.set(price_str)
                self.logger.info("Removed dollar sign from discount price entry")

            # Store the rounded numerical value
            try:
                self.initial_discount_price = round(float(price_str), 2)
                self.logger.info(f"Stored initial discount price: {self.initial_discount_price}")
            except ValueError:
                self.initial_discount_price = None
                self.logger.error("Invalid discount price format, unable to store initial value")
        except Exception as e:
            self.logger.error(f"Error handling discount price focus in: {e}")

    def on_discount_price_focus_out(self, event=None):
        """Adds '$' symbol to the discount price when focus is lost."""
        self.logger.info("Handling focus out event for discount price entry")

        try:
            price_str = self.discount_var.get()

            # Check if the price string is empty or invalid, set it to '$0'
            if not price_str or not price_str.replace('$', '').strip().replace('.', '', 1).isdigit():
                self.discount_var.set("$0")
                final_discount_price = 0.0
                self.logger.info("Invalid or empty discount price, set to $0")
            else:
                if not price_str.startswith('$'):
                    self.discount_var.set(f"${price_str}")
                    self.logger.info(f"Added dollar sign to discount price: {price_str}")

                try:
                    final_discount_price = round(float(price_str.lstrip('$')), 2)
                except ValueError:
                    final_discount_price = None
                    self.logger.error("Invalid discount price format, unable to process")

            # Trigger discount calculation only if the price has changed
            if self.initial_discount_price != final_discount_price:
                self.last_changed = 'price'
                self.calculate_discount('price')  # Pass 'price' as the argument
                self.logger.info("Discount price changed, recalculating discount")
            else:
                # Optionally, handle the case where the price hasn't changed
                self.logger.info("Discount price unchanged, no recalculation needed")
        except Exception as e:
            self.logger.error(f"Error handling discount price focus out: {e}")

    def on_percentage_changed(self, *args):
        self.last_changed = 'percentage'
        self.calculate_discount()

    def on_discount_percentage_focus_in(self, event=None):
        self.logger.info("Handling focus in event for discount percentage entry")

        try:
            percentage_str = self.percent_discount_var.get()

            if percentage_str.endswith('%'):
                percentage_str = percentage_str.rstrip('%')
                self.percent_discount_var.set(percentage_str)
                self.logger.info("Removed percentage sign from discount percentage entry")

            # Now try converting the stripped string to a float
            try:
                self.initial_percent_discount = round(float(percentage_str), 2)
                self.logger.info(f"Stored initial discount percentage: {self.initial_percent_discount}")
            except ValueError:
                self.initial_percent_discount = None
                self.logger.error("Invalid discount percentage format, unable to store initial value")
        except Exception as e:
            self.logger.error(f"Error handling discount percentage focus in: {e}")

    def on_discount_percentage_focus_out(self, event=None):
        self.logger.info("Handling focus out event for discount percentage entry")

        try:
            percentage_str = self.percent_discount_var.get()

            # Check if the percentage string is empty or invalid, set it to '0%'
            if not percentage_str or not percentage_str.replace('%', '').strip().isdigit():
                self.percent_discount_var.set("0%")
                final_percent_discount = 0.0
                self.logger.info("Invalid or empty discount percentage, set to 0%")
            else:
                if not percentage_str.endswith('%'):
                    self.percent_discount_var.set(f"{percentage_str}%")
                    self.logger.info(f"Added percentage sign to discount percentage: {percentage_str}")

                try:
                    final_percent_discount = round(float(percentage_str.strip('%')), 2)
                except ValueError:
                    final_percent_discount = None
                    self.logger.error("Invalid discount percentage format, unable to process")

            # Trigger discount calculation only if the percentage has changed
            if self.initial_percent_discount != final_percent_discount:
                self.last_changed = 'percentage'
                self.calculate_discount('percentage')  # Pass 'percentage' as the argument
                self.logger.info("Discount percentage changed, recalculating discount")
            else:
                # Optionally, handle the case where the percentage hasn't changed
                self.logger.info("Discount percentage unchanged, no recalculation needed")
        except Exception as e:
            self.logger.error(f"Error handling discount percentage focus out: {e}")

    def custom_float_format(self, value):
        """Formats the float value to string with two decimal places."""
        return "{:.2f}".format(value)

    def calculate_discount(self, based_on):
        self.logger.info(f"Calculating discount based on: {based_on}")

        try:
            price_str = self.regular_product_price_var.get().lstrip('$')
            price = Decimal(price_str) if price_str else Decimal('0')

            if based_on == 'percentage':
                percentage_str = self.percent_discount_var.get().strip('%')
                percentage = Decimal(percentage_str) if percentage_str else Decimal('0')
                discount_price = (price * percentage / Decimal('100')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                self.discount_var.set(f"${discount_price:.2f}")
                self.logger.info(f"Discount calculated based on percentage: ${discount_price:.2f}")

            elif based_on == 'price':
                discount_str = self.discount_var.get().strip('$')
                discount = Decimal(discount_str) if discount_str else Decimal('0')
                percentage = ((discount / price) * Decimal('100')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                
                # Adjust the format to allow for decimal percentages
                formatted_percentage = "{:.2f}%".format(percentage)
                self.percent_discount_var.set(formatted_percentage)
                self.logger.info(f"Discount calculated based on price: {formatted_percentage}")
            self.calculate_discount_fields()
        except (ValueError, InvalidOperation) as e:
            self.logger.error(f"Error calculating discount: {e}")

    def calculate_discount_fields(self):
        """
        Calculate the fields related to discounts such as the discounted product price, 
        IVU tax after discount, and the total price after discount including IVU tax.
        """
        self.logger.info("Calculating discount fields")
        # Helper function to strip characters and convert to Decimal
        def clean_and_convert(value, strip_char=None):
            if strip_char:
                value = value.replace(strip_char, '')
            try:
                return Decimal(value)
            except (ValueError, InvalidOperation):
                self.logger.error(f"Invalid format for value '{value}'. Setting to Decimal('0')")
                return Decimal('0')

        # Get values and clean them
        product_price_plus_ivu = clean_and_convert(self.product_price_plus_ivu_var.get(), '$')
        discount_price = clean_and_convert(self.discount_var.get(), '$')
        discount_percentage = clean_and_convert(self.percent_discount_var.get(), '%')

        # Define the tax rate
        tax_rate = Decimal('0.115')

        # Correctly calculate the original product price and IVU Tax
        original_product_price = (product_price_plus_ivu / (1 + tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        ivu_tax = (original_product_price * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Check if there's a discount
        has_discount = discount_price > 0 or discount_percentage > 0
        # Calculate Discounted Prices
        try:
            if has_discount:
                # Determine the discount amount
                if discount_price > 0:
                    discount_amount = discount_price
                else:
                    discount_amount = (original_product_price * (discount_percentage / Decimal('100'))).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

                # Apply the discount to the original product price
                product_price_after_discount = (original_product_price - discount_amount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

                # Recalculate the IVU tax based on the discounted price
                ivu_tax_after_discount = (product_price_after_discount * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

                # Calculate the total price after discount including IVU tax
                product_price_plus_ivu_discount = (product_price_after_discount + ivu_tax_after_discount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
            else:
                product_price_after_discount = original_product_price
                ivu_tax_after_discount = ivu_tax
                product_price_plus_ivu_discount = product_price_plus_ivu

            # Update the fields
            self.product_price_after_discount_var.set(f"${product_price_after_discount:.2f}")
            self.ivu_tax_after_discount_var.set(f"${ivu_tax_after_discount:.2f}")
            self.product_price_minus_discount_plus_ivu_var.set(f"${product_price_plus_ivu_discount:.2f}")
            self.logger.info("Discount fields successfully calculated and updated")
        except Exception as e:
            self.logger.error(f"Error calculating discount fields: {e}")

    def recalculate_original_price_and_tax(self):
        """
        Recalculates the original product price and IVU tax based on the product price
        including IVU (price_plus_ivu).
        """
        self.logger.info("Recalculating original product price and IVU tax")

        # Extract and clean the product price (+ IVU) value
        price_plus_ivu_str = self.product_price_plus_ivu_var.get().lstrip('$')
        try:
            price_plus_ivu = Decimal(price_plus_ivu_str)
        except ValueError:
            price_plus_ivu = Decimal('0')
            self.logger.error("Invalid format for product price plus IVU, using Decimal('0')")

        # Define the tax rate
        tax_rate = Decimal('0.115')

        # Calculate the original product price before tax
        original_product_price = (price_plus_ivu / (1 + tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Calculate the IVU tax based on the original product price
        IVU_tax = (original_product_price * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        price_plus_ivu_str = original_product_price + IVU_tax
        # Update the IVU tax and original product price fields
        self.ivu_tax_var.set(f"${IVU_tax:.2f}")
        self.regular_product_price_var.set(f"${original_product_price:.2f}")
        self.product_price_plus_ivu_var.set(f"${price_plus_ivu_str:.2f}")
        self.logger.info("Original product price and IVU tax recalculated and updated")

    def pick_date(self):
        """
        Opens a calendar widget to pick a date. The selected date is then formatted
        and inserted into the sold_date_entry field.
        """
        self.logger.info("Opening calendar widget to pick a date")
        def grab_date():
            try:
                selected_date = cal.selection_get()  # Get the selected date
                formatted_date = selected_date.strftime('%m/%d/%Y')  # Format the date

                self.sold_date_entry.config(state="normal")  # Enable the entry widget
                self.sold_date_entry.delete(0, tk.END)  # Clear the entry field
                self.sold_date_entry.insert(0, formatted_date)  # Insert the formatted date
                self.sold_date_entry.config(state="disabled")  # Disable the entry widget

                top.destroy()  # Close the Toplevel window
                self.logger.info(f"Date picked and set: {formatted_date}")
            except Exception as e:
                self.logger.error(f"Error in grabbing date from calendar: {e}")

        def select_today_and_close(event):
            try:
                cal.selection_set(datetime.today())  # Set selection to today's date
                grab_date()  # Then grab the date and close
                self.logger.info("Today's date selected and set")
            except Exception as e:
                self.logger.error(f"Error in selecting today's date and closing calendar: {e}")

        top = tk.Toplevel(self)
        today = datetime.today()
        cal = Calendar(top, selectmode='day', year=today.year, month=today.month, day=today.day)
        cal.pack(pady=20)    # Set focus to the Toplevel window and bind the Enter key
        top.focus_set()
        top.bind('<Return>', select_today_and_close)
        cal.bind("<<CalendarSelected>>", lambda event: grab_date())
    
    def clear_date(self):
        """
        Clears the date from the sold_date_entry widget. The widget is temporarily 
        enabled for clearing and then disabled again to prevent manual inputs.
        """
        self.logger.info("Clearing the date from the sold date entry widget")

        try:
            self.sold_date_entry.config(state="normal")  # Enable the entry widget
            self.sold_date_entry.delete(0, tk.END)  # Clear the entry field
            self.sold_date_entry.config(state="disabled")  # Disable the entry widget
            self.logger.info("Sold date entry cleared")
        except Exception as e:
            self.logger.error(f"Error in clearing the sold date entry: {e}")

    def open_hyperlink(self, event):
        """
        Opens the hyperlink URL in a web browser when clicked. This function is
        triggered by a mouse click event on a hyperlink in the order_link_text widget.
        """
        self.logger.info("Attempting to open hyperlink from the text widget")

        try:
            start_index = self.order_link_text.index("@%s,%s" % (event.x, event.y))
            tag_indices = list(self.order_link_text.tag_ranges('hyperlink'))
            for start, end in zip(tag_indices[0::2], tag_indices[1::2]):
                if self.order_link_text.compare(start_index, ">=", start) and self.order_link_text.compare(start_index, "<=", end):
                    url = self.order_link_text.get(start, end)
                    webbrowser.open(url)
                    self.logger.info(f"Opened hyperlink: {url}")
                    return "break"
        except Exception as e:
            self.logger.error(f"Error when opening hyperlink: {e}")


    def back_to_main(self):
        """
        Closes the settings window and returns to the main application window.
        It also reloads the settings and refreshes the folder list based on any updates.
        """
        self.logger.info("Returning to the main window from settings")

        self.settings_window.destroy()
        self.master.deiconify()
        self.master.state('zoomed')
        
        self.load_settings()
        self.combine_and_display_folders()

        self.logger.info("Returned to the main window and updated settings and folders list")

    def choose_inventory_folder(self):
        """
        Opens a file dialog to select the inventory folder. Updates the inventory folder path in the 
        application, the corresponding label in the settings, and saves the new setting.
        """
        self.logger.info("Choosing inventory folder")

        # Open a dialog to choose the inventory folder
        inventory_folder = filedialog.askdirectory()
        if inventory_folder:
            self.inventory_folder = inventory_folder
            self.inventory_folder_label.config(text=inventory_folder)  # Update the label directly
            self.save_settings()  # Save the updated settings
            self.combine_and_display_folders()  # Refresh the folders list based on new settings

        self.logger.info("Inventory folder chosen and settings updated")

    @staticmethod
    def custom_sort_key(s):
        """
        Defines a custom sorting key for sorting strings (e.g., folder names). 
        Sorts based on the length of the first word, followed by alphanumeric sorting of the words.
        This sorting is case-insensitive and numbers will sort naturally before letters.
        """
        # A regular expression to match words in the folder name.
        # Words are defined as sequences of alphanumeric characters and underscores.
        words = re.findall(r'\w+', s.lower())
        
        # The key will be a tuple consisting of the length of the first word,
        # the first word itself (for alphanumeric sorting), and then the rest of the words.
        # Lowercase all words for case-insensitive comparison, numbers will sort naturally before letters.
        return (len(words[0]),) + tuple(words)


    def choose_sold_folder(self):
        """
        Opens a file dialog to select the sold inventory folder. Updates the sold folder path in the 
        application, the corresponding label in the settings, saves the new setting, and updates the 
        database with the new sold folder path.
        """
        self.logger.info("Choosing sold inventory folder")

        # Open a dialog to choose the sold folder
        self.sold_folder = filedialog.askdirectory()
        if self.sold_folder:
            self.sold_folder_label.config(text=self.sold_folder)  # Update the label directly
            self.save_settings()  # Save the updated settings
            self.logger.info("Sold inventory folder chosen and settings updated")

        # Update the Sold Folder path in the database
        try:
            self.db_manager.cur.execute('''
                INSERT INTO folder_paths (Folder, Path) VALUES ('Sold', ?)
                ON CONFLICT(Folder) DO UPDATE SET Path = excluded.Path;
            ''', (self.sold_folder,))
            self.db_manager.conn.commit()
            self.logger.info("Sold folder path updated in the database")
        except Exception as e:
            self.logger.error(f"Error updating sold folder path in database: {e}")

    def choose_to_sell_folder(self):
        """
        Opens a file dialog to select the 'Products to Sell' folder. Updates the 'to sell' folder path 
        in the application and the corresponding label in the settings, and saves the new setting.
        """
        self.logger.info("Choosing 'Products to Sell' folder")

        # Open a dialog to choose the 'to sell' folder
        self.to_sell_folder = filedialog.askdirectory()
        if self.to_sell_folder:
            self.to_sell_folder_label.config(text=self.to_sell_folder)  # Update the label directly
            self.save_settings()  # Save the updated settings
            self.logger.info("'Products to Sell' folder chosen and settings updated")

    def save_settings(self):
        """
        Gathers the paths for inventory, sold, and to sell folders and writes them to a 'folders_paths.txt' file. 
        This method saves the current folder settings persistently.
        """
        self.logger.info("Saving folder settings to 'folders_paths.txt'")

        try:
            with open("folders_paths.txt", "w") as file:
                file.write(f"{self.inventory_folder}\n{self.sold_folder}\n{self.to_sell_folder}")
            self.logger.info("Folder settings successfully written to 'folders_paths.txt'")
        except Exception as e:
            self.logger.error(f"Error saving folder settings to file: {e}")


    def load_and_display_image(self, current_row_num, product_image_col_num, product_id):
        """
        Loads and displays the product image in a separate thread. The image is fetched from the 
        Excel workbook based on the row and column numbers. If the image is not in the cache, it 
        is loaded from the workbook, cached, and then displayed.
        """
        # Before starting the thread for loading the image
        self.logger.info(f"Starting thread to load image for product ID: {product_id}")

        def task():
            self.logger.info(f"Starting image loading task: Row {current_row_num}, Column {product_image_col_num}")

            if not self.running or self.current_product_id != product_id:
                self.logger.info("Task exited: Application no longer running or product changed")
                return
            wb = None
            try:
                image_data = self.get_image_data(current_row_num, product_image_col_num)
                if not image_data:
                    # Image not in cache, load from workbook and cache it
                    wb = self.load_workbook_cached(self.excel_manager.filepath)
                    sheet = wb[self.excel_manager.sheet_name]
                    for image in sheet._images:
                        if image.anchor._from.row == current_row_num and image.anchor._from.col == product_image_col_num:
                            image_data = image._data()
                            self.image_cache[(current_row_num, product_image_col_num)] = image_data
                            # Log when image is found and cached
                            self.logger.info("Image found and cached")
                            break
                    wb.close()

                if image_data:
                    # Load image from cached data
                    with io.BytesIO(image_data) as image_stream:
                        pil_image = Image.open(image_stream)
                        # Resize the image using PIL
                        desired_size = (100, 100)
                        resized_image = pil_image.resize(desired_size)
                        # Convert the resized image to Tkinter PhotoImage

                    if self.running and self.current_product_id == product_id:
                        self.logger.info("Scheduling image update in main thread")
                        self.after(0, lambda: self.update_image_label(resized_image))
                    else:
                        self.logger.error("Skipped image update: Application no longer running or product changed")
                else:
                    self.logger.error("Image not found in workbook or cache")
                    if self.running and self.current_product_id == product_id:
                        self.after(0, lambda: self.product_image_label.config(text="Product image not found"))

            except Exception as e:
                self.logger.error(f"Error loading image: {e}")
                if self.running and self.current_product_id == product_id:
                    self.after(0, lambda: self.product_image_label.config(text="Error loading image"))
            finally:
                if wb:
                    wb.close()
                    self.logger.info("Workbook closed in thread")

        threading.Thread(target=task).start()

    def load_workbook_cached(self, path):
        """
        Loads an Excel workbook from the given path with caching. 
        If the workbook at the same path is already loaded, it uses the cached version instead 
        of reloading it. This improves performance by avoiding redundant loading of the same workbook.
        """
        self.logger.info(f"Loading workbook from path: {path}")

        # Check if the path is different from the cached path or the cache is None
        if path != self.workbook_path or self.workbook_cache is None:
            # Load the workbook and update the cache
            try:
                self.workbook_cache = openpyxl.load_workbook(path, data_only=True)
                self.workbook_path = path
                self.logger.info("Workbook loaded and cached")
            except Exception as e:
                self.logger.error(f"Error loading workbook from path {path}: {e}")
                raise

        return self.workbook_cache

    def cache_images(self, workbook_path, sheet_name):
        """
        Loads images from a specified Excel sheet and caches them. 
        Each image is associated with its cell position (row and column) in the sheet.
        """
        self.logger.info(f"Caching images from workbook '{workbook_path}', sheet '{sheet_name}'")

        wb = None
        try:
            wb = openpyxl.load_workbook(workbook_path, data_only=True)
            sheet = wb[sheet_name]

            for image in sheet._images:
                row, col = image.anchor._from.row, image.anchor._from.col
                key = (row, col)
                self.logger.info(f"Caching image at row {row}, column {col}")
                self.image_cache[key] = image._data()

            self.logger.info("Finished caching images")
        except Exception as e:
            self.logger.error(f"Error caching images from workbook: {e}")
        finally:
            if wb:
                wb.close()
                self.logger.info("Workbook closed after caching images")

    def get_image_data(self, row, col):
        key = (row, col)
        return self.image_cache.get(key, None)

    def update_image_label(self, pil_image):
        if self.running:
            self.logger.info("Updating image label in main thread")
            tk_photo = ImageTk.PhotoImage(pil_image)
            self.product_image_label.config(image=tk_photo)
            self.product_image_label.image = tk_photo  # Keep a reference
        else:
            self.logger.error("Skipped updating image label: Application no longer running")

    def open_product_folder(self, folder_path):
        if sys.platform == "win32":
            os.startfile(folder_path)
        elif sys.platform == "darwin":  # macOS
            subprocess.run(["open", folder_path])
        else:  # Linux variants
            subprocess.run(["xdg-open", folder_path])

    def excel_value_to_bool(self, value):
        """
        Converts an Excel cell value to a boolean. The conversion considers string values 
        like 'yes', 'true', '1', and numeric values, treating NaN and unrecognized formats as False.
        """

        # Before checking the value
        self.logger.info(f"Converting Excel value to boolean: {value}")

        if pd.isnull(value):
            return False

        if isinstance(value, str):
            result = value.strip().lower() in ['yes', 'true', '1']
            self.logger.info(f"Converted string '{value}' to boolean: {result}")
            return result
        elif isinstance(value, (int, float)):
            result = bool(value)
            self.logger.info(f"Converted numeric value '{value}' to boolean: {result}")
            return result

        # Log the default case when the value doesn't match expected types or formats
        self.logger.info("Value format unrecognized, defaulting to False")
        return False

    def update_to_sell_after_color(self):
        """
        Updates the background color of the 'to sell after' label based on the date comparison. 
        If the date has passed or is today, the background is set to green, otherwise, it's set to white.
        """

        # Log before starting the process
        self.logger.info("Updating 'To Sell After' label color based on date comparison")

        # Get today's date
        today = date.today()

        # Get the date from the to_sell_after_var entry
        to_sell_after_str = self.to_sell_after_var.get()
        if to_sell_after_str:
            try:
                # Parse the date string to a date object
                to_sell_after_date = datetime.strptime(to_sell_after_str, "%m/%d/%Y").date()
                self.logger.info(f"'To Sell After' date parsed: {to_sell_after_date}")

                # If the to_sell_after date is today or has passed, change the label's background color to green
                if to_sell_after_date <= today:
                    self.to_sell_after_label.config(background='light green')
                    self.logger.info("'To Sell After' label background set to green")
                else:
                    self.to_sell_after_label.config(background='white')
                    self.logger.info("'To Sell After' label background reset to white")
            except ValueError:
                # If there's a ValueError, it means the string was not in the expected format
                # Handle incorrect date format or clear the background
                self.to_sell_after_label.config(background='white')
                self.logger.error("Incorrect date format in 'To Sell After', resetting background color")
    
    def checkbox_click_control(self, var):
        """
        Controls the checkbox click event based on the current edit mode. 
        If not in edit mode, it prevents changing the checkbox's state.
        """

        # Log before checking the edit mode
        self.logger.info(f"Checkbox click control invoked, edit mode: {self.edit_mode}")

        if not self.edit_mode:
            # Log preventing checkbox state change
            self.logger.info("Checkbox state change prevented due to non-edit mode")
            return "break"  # Stop the event from propagating further

        # Log allowing checkbox state change
        self.logger.info("In edit mode, allowing checkbox state change")

    def toggle_edit_mode(self):
        """
        Toggles the edit mode of the application. When in edit mode, certain widgets are made editable, 
        and key bindings are set for saving and editing. When not in edit mode, the widgets are set to 
        non-editable, and key bindings are removed.
        """

        # Log before toggling the edit mode
        self.logger.info("Toggling edit mode")

        self.edit_mode = not self.edit_mode
        state = 'normal' if self.edit_mode else 'disabled' 
        readonly_state = 'readonly' if self.edit_mode else 'disabled'

        # Log the state of the edit mode after toggling
        self.logger.info(f"Edit mode set to: {self.edit_mode}")   
        
        self.order_date_entry.config(state='disabled')
        self.sold_date_button.config(state=state)
        self.clear_button.config(state=state)       
        self.to_sell_after_entry.config(state='disabled')
        self.payment_type_combobox.config(state=readonly_state)
        self.asin_entry.config(state=state)
        self.product_id_entry.config(state='disabled')
        self.rack_id_entry.config(state=state)
        self.product_name_text.config(state='disabled')
        self.fair_market_value_entry.config(state='disabled')
        self.regular_product_price_entry.config(state='disabled')
        self.ivu_tax_entry.config(state='disabled')
        self.product_price_plus_ivu_entry.config(state=state)
        self.discount_entry.config(state=state)
        self.percent_discount_entry.config(state=state)
        self.sold_price_entry.config(state=state)
        self.save_button.config(state=state)
        self.comments_text.config(state=state)
        self.product_description_text.config(state=state)

        if self.edit_mode:
            self.logger.info("Edit mode enabled, setting widget states and bindings")

            self.product_name_text.bind("<Button-1>", lambda e: None)
            self.master.bind('<Return>', self.save_on_key_handler)
            self.master.bind('<Escape>', self.edit_on_key_handler)
        else:
            self.logger.info("Edit mode disabled, resetting widget states and unbinding keys")
            self.product_name_text.bind("<Button-1>", lambda e: "break")

            self.master.unbind('<Return>')
            self.master.unbind('<Escape>')
            self.master.bind('<Return>', self.edit_on_key_handler)

    def save(self):
        """
        Saves the updated product information from the form into the Excel file and moves the 
        product folder to the appropriate location based on the current status (sold, damaged, personal, etc.).
        """

        # Log before starting the save process
        self.logger.info("Saving product information")
        # Extract values from the widgets
        sold_price = self.sold_price_entry.get()
        sold_date = self.sold_date_var.get()  # Assuming it's a StringVar associated with an Entry
        payment_type = self.payment_type_var.get()  # Similarly, for payment type

        # Check if any of the fields have data
        if sold_price or sold_date or payment_type:
            # Check if all required fields are filled
            if not (sold_price and sold_date and payment_type):
                messagebox.showwarning("Incomplete Data", "Please fill in Sold Price, Sold Date, and Payment Type.")
                self.logger.error("Incomplete data for saving")
                return  # Return without saving

        # Update the 'Sold' checkbox based on the 'Sold Date' entry
        if self.sold_date_var.get():
            # If 'Sold Date' is not empty, check 'Sold'
            self.sold_var.set(True)
        else:
            # If 'Sold Date' is empty, uncheck 'Sold'
            self.sold_var.set(False)

        def to_float(value):
            try:
                # Remove any non-numeric characters like $ and %, then convert to float
                numeric_value = value.replace('$', '').replace('%', '')
                return float(numeric_value)
            except ValueError:
                # Return the original value if it can't be converted
                return value
            
        def remove_dollar_sign(value):
            return value.replace('$', '') if isinstance(value, str) else value
        
        product_id = self.product_id_var.get().strip().upper()

        # Ensure that the Excel file path and sheet name are set.
        filepath, sheet_name = self.load_excel_path_and_sheet()

        if not filepath or not sheet_name:
            messagebox.showerror("Error", "Excel file path or sheet name is not set.")
            return

        # Collect the data from the form.
        product_data = {
            'Cancelled Order': self.cancelled_order_var.get(),
            'Damaged': self.damaged_var.get(),
            'Personal': self.personal_var.get(),
            'Reviewed': self.reviewed_var.get(),
            'Pictures Downloaded': self.pictures_downloaded_var.get(),
            'Uploaded to Site': self.uploaded_to_site_var.get(),
            'Sold': self.sold_var.get(),

            'To Sell After': self.to_sell_after_var.get(),
            'Product Name': self.product_name_text.get("1.0", tk.END).strip(),
            'Rack ID': self.rack_id_var.get(),
            'Sold Price': self.sold_price_var.get(),
            'Payment Type': self.payment_type_var.get(),
            'Sold Date': self.sold_date_var.get(),
            'Comments': self.comments_text.get("1.0", tk.END).strip(),
            'Product Description': self.product_description_text.get("1.0", tk.END).strip(),
            'Fair Market Value': to_float(remove_dollar_sign(self.fair_market_value_var.get())),
            'Discount': to_float(remove_dollar_sign(self.discount_var.get())),
            'Discount Percentage': to_float(remove_dollar_sign(self.percent_discount_var.get())),
            'Product Price': to_float(remove_dollar_sign(self.regular_product_price_var.get())),
            'IVU Tax': to_float(remove_dollar_sign(self.ivu_tax_var.get())),
            'Product Price After IVU': to_float(remove_dollar_sign(self.product_price_plus_ivu_var.get())),

            'Product Price After Discount': to_float(remove_dollar_sign(self.product_price_after_discount_var.get())),
            'IVU Tax After Discount': to_float(remove_dollar_sign(self.ivu_tax_after_discount_var.get())),
            'Product Price After IVU and Discount': to_float(remove_dollar_sign(self.product_price_minus_discount_plus_ivu_var.get())),

            'Sold Price': to_float(remove_dollar_sign(self.sold_price_var.get())),
            # ... and so on for the rest of your form fields.
        }

        # Use the ExcelManager method to save the data.
        try:
            self.logger.info("Attempting to save data to Excel")
            self.excel_manager.save_product_info(product_id, product_data)
            messagebox.showinfo("Success", "Product information updated successfully.")
            self.logger.info("Product information updated successfully in Excel")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save changes to Excel file: {e}")
            self.logger.error(f"Failed to save changes to Excel file: {e}")
            return
        
        # Folder movement logic
        current_folder_path = self.get_folder_path_from_db(product_id)
        if not current_folder_path:
            messagebox.showerror("Error", f"No current folder path found for Product ID {product_id}")
            return
        
        folder_name = os.path.basename(current_folder_path)

        # Initialize variables for folder paths
        damaged_folder_path = os.path.join(os.path.dirname(self.inventory_folder), "Damaged")
        personal_folder_path = os.path.join(os.path.dirname(self.inventory_folder), "Personal")

        # Create Damaged and Personal folders if they do not exist
        for folder in [damaged_folder_path, personal_folder_path]:
            if not os.path.exists(folder):
                os.makedirs(folder)

        product_name = self.product_name_text.get ("1.0", tk.END).strip()


        # Decide target folder based on checkbox statuses and other conditions
        if self.sold_var.get():
            target_folder_path = self.sold_folder
        elif self.damaged_var.get():
            target_folder_path = damaged_folder_path
        elif self.personal_var.get():
            target_folder_path = personal_folder_path
        else:
            to_sell_after_str = self.to_sell_after_var.get()
            try:
                to_sell_after_date = datetime.strptime(to_sell_after_str, "%m/%d/%Y").date() if to_sell_after_str else None
            except ValueError as e:
                messagebox.showerror("Error", f"Invalid 'To Sell After' date format: {e}")
                return

            today = date.today()
            if to_sell_after_date and to_sell_after_date <= today:
                target_folder_path = self.to_sell_folder
            else:
                target_folder_path = self.inventory_folder
        # Use #print statements to debug the current and target folder paths
        #print(f"Current folder path: {current_folder_path}")
        #print(f"Target folder path: {target_folder_path}")

        # Check if the target folder is determined and it's not the same as the current folder
        if target_folder_path and os.path.isdir(current_folder_path) and current_folder_path != target_folder_path:
            try:
                # Perform the move operation
                new_folder_path = self.move_product_folder(current_folder_path, folder_name, target_folder_path, product_name)
                # folder_name is not defined. not sure how to get the folder_name
                
                new_folder_name = os.path.basename(new_folder_path).strip()  # Extract folder name from the path

                self.db_manager.delete_folder_path(folder_name)

                # Save the new folder path in the database
                self.db_manager.save_folder_path(new_folder_name, new_folder_path)
                
                messagebox.showinfo("Folder Moved", f"Folder for '{product_id}' moved successfully to the new location.")
                #print(f"Folder for '{product_id}' moved from {current_folder_path} to {new_folder_path}")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to move the folder: {e}")
        


        #self.update_folder_path_and_name(product_id, product_name, current_folder_path, target_folder_path)

        self.refresh_and_select_product(product_id)
        self.logger.info("Product information saved and folder moved successfully")
        doc_data = (product_id, product_id, self.product_name_var.get())  # Construct the doc_data tuple
        self.create_word_doc(doc_data, iid="dummy", show_message=True)  # Call create_word_doc with dummy iid
        self.toggle_edit_mode()



        # Unbind the Enter and Escape keys
        #self.master.unbind('<Return>')
        #self.master.unbind('<Escape>')
    
        # Additionally, you might want to re-bind the Enter key to the edit_button's command
        # if you want to be able to press Enter to switch to edit mode again
        #self.master.bind('<Return>', lambda e: self.edit_on_key_handler.invoke())

    def get_folder_path_from_db(self, product_id):
        """
        Retrieves the folder path for a given product ID from the database. 
        The function assumes that the folder name in the database starts with the product ID followed by a space.
        """

        # Log before executing the database query
        self.logger.info(f"Fetching folder path for product ID: {product_id} from the database")

        self.db_manager.cur.execute("SELECT Path FROM folder_paths WHERE Folder LIKE ?", (product_id + ' %',))
        result = self.db_manager.cur.fetchone()


        return result[0] if result else None


    def get_folder_names_from_db(self):
        """
        Retrieves all folder names from the database and returns them as a list. 
        This function queries the 'folder_paths' table to get the names of folders.
        """

        # Log before executing the database query
        self.logger.info("Fetching folder names from the database")

        self.db_manager.cur.execute("SELECT Folder FROM folder_paths")
        folder_names = [row[0] for row in self.db_manager.cur.fetchall()]

        # Log after successfully fetching the data
        self.logger.info("Successfully fetched folder names from the database")

        return folder_names


    def select_excel_database(self):
        """
        Opens a file dialog for the user to select an Excel database file. 
        It updates the ExcelManager instance with the selected file path and sheet name, 
        saves these settings, and loads the data from the selected Excel file.
        """

        # Log before opening the file dialog
        self.logger.info("Opening file dialog to select Excel database")

        filepath = filedialog.askopenfilename(
            title="Select Excel Database",
            filetypes=[("Excel Files", "*.xlsx *.xls *.xlsm")]
        )
        if filepath:
            self.excel_manager.filepath = filepath  # Save the filepath to the ExcelManager instance
            self.logger.info(f"Excel database selected: {filepath}")
            xls = pd.ExcelFile(filepath)
            sheet_names = xls.sheet_names
            if sheet_names:
                # Automatically select the first sheet if available
                self.excel_manager.sheet_name = sheet_names[0]  # Save the sheet name to the ExcelManager instance
                self.save_excel_settings(filepath, sheet_names[0])  # Save settings
                self.excel_manager.load_data()  # Load the data
                self.update_excel_label()  # Update the label
                self.logger.info(f"Excel sheet selected and data loaded: {sheet_names[0]}")
        xls = pd.ExcelFile(filepath) # delete ?
        sheet_names = xls.sheet_names # delete ?
        self.ask_sheet_name(sheet_names, filepath)  # Pass filepath here
        self.logger.info("Asked for sheet name selection")

    def update_excel_label(self):
        """
        Updates the label in the GUI to display the current Excel file path and the selected sheet name.
        """

        # Log before updating the label
        self.logger.info("Updating Excel database label in the GUI")

        excel_db_text = f"{self.excel_manager.filepath} - Sheet: {self.excel_manager.sheet_name}"
        self.excel_db_label.config(text=excel_db_text)

        # Log after successfully updating the label
        self.logger.info("Excel database label updated")

    def ask_sheet_name(self, sheet_names, filepath):
        """
        Opens a new window for the user to select a sheet from the given Excel file. 
        The window displays a list of sheet names and allows the user to make a selection.
        """

        # Log before opening the sheet selection window
        self.logger.info("Opening window to select a sheet from the Excel file")

        sheet_window = tk.Toplevel(self)
        sheet_window.title("Select a Sheet")

        listbox = tk.Listbox(sheet_window, exportselection=False)
        listbox.pack(padx=10, pady=10)

        # Populate listbox with sheet names
        for sheet in sheet_names:
            listbox.insert('end', sheet)

        # Set the default selection
        default_sheet_index = sheet_names.index(self.default_sheet) if self.default_sheet in sheet_names else 0
        listbox.selection_set(default_sheet_index)
        listbox.activate(default_sheet_index)

        # Bind double-click event to the listbox
        listbox.bind('<Double-1>', lambda event: self.confirm_sheet_selection(event, listbox, filepath))

        confirm_button = ttk.Button(sheet_window, text="Confirm", command=lambda: self.confirm_sheet_selection(None, listbox, filepath))
        confirm_button.pack(pady=(0, 10))

        # Log after setting up the sheet selection window
        self.logger.info("Sheet selection window set up")
        sheet_window.wait_window()
        # Log after the sheet selection window is closed
        self.logger.info("Sheet selection window closed")

    def confirm_sheet_selection(self, event, listbox, filepath):
        """
        Confirms the user's sheet selection from the listbox and updates the ExcelManager instance 
        with the selected sheet. Closes the sheet selection window after the selection is confirmed.
        """

        # Log before confirming sheet selection
        self.logger.info("Confirming sheet selection from the listbox")

        selection_index = listbox.curselection()
        if selection_index:
            selected_sheet = listbox.get(selection_index[0])
            # Log the selected sheet
            self.logger.info(f"Selected sheet: {selected_sheet}")

            self.select_excel_sheet(selected_sheet, filepath)
            listbox.master.destroy()  # Closes the sheet_window

            # Log after successfully selecting the sheet and closing the window
            self.logger.info("Sheet selected and sheet selection window closed")

    def select_excel_sheet(self, selected_sheet, filepath):
        """
        Updates the ExcelManager with the newly selected sheet and file path. 
        Loads data from the selected sheet and updates the Excel file and sheet label in the GUI.
        """
        # Log before updating ExcelManager with the new sheet
        self.logger.info(f"Selecting Excel sheet: {selected_sheet}")

        self.excel_manager.filepath = filepath
        self.excel_manager.sheet_name = selected_sheet
        self.excel_manager.load_data()
        self.update_excel_label()
        self.save_excel_settings(filepath, selected_sheet)

        # Log after successfully updating the ExcelManager and saving settings
        self.logger.info("Excel sheet selected and data loaded")

    def save_excel_settings(self, filepath, sheet_name):
        """
        Saves the current Excel file path and sheet name to a text file for persistence.
        """
        # Log before attempting to save settings
        self.logger.info("Saving Excel settings to file")

        try:
            with open('excel_and_sheet_path.txt', 'w') as f:
                f.write(f"{filepath}\n{sheet_name}")
            self.update_excel_label()  # Update the label when settings are saved
            self.logger.info("Excel settings saved successfully")
        except Exception as e:
            messagebox.showerror("Error", f"Unable to save settings: {str(e)}")
            self.logger.error(f"Failed to save Excel settings: {e}")

    def load_excel_path_and_sheet(self):
        """
        Loads the saved Excel file path and sheet name from a text file. 
        If the settings file is not found or an error occurs, it returns None for both filepath and sheet_name.
        """

        # Log before attempting to load Excel settings
        self.logger.info("Loading Excel settings from file")

        try:
            with open('excel_and_sheet_path.txt', 'r') as f:
                filepath, sheet_name = f.read().strip().split('\n', 1)
                self.logger.info("Excel settings loaded successfully")
                return filepath, sheet_name
        except FileNotFoundError:
            self.logger.error("Excel settings file not found")
            return None, None
        except Exception as e:
            messagebox.showerror("Error", f"Unable to load settings: {str(e)}")
            self.logger.error(f"Failed to load Excel settings: {e}")
            return None, None

    # def update_links_in_excel(self):
    #     """
    #     Updates the 'Order Link' column in the Excel file based on hyperlinks in the 'Product Name' column. 
    #     This function assumes the existence of a settings file with Excel path and sheet name, 
    #     and the specific structure of the Excel sheet.
    #     """

    #     # Log the start of the link update process
    #     self.logger.info("Starting the process to update links in the Excel file")

    #     try:
    #         with open('excel_and_sheet_path.txt', 'r') as file:
    #             lines = file.readlines()
    #             excel_path = lines[0].strip()
    #             sheet_name = lines[1].strip()

    #         workbook = openpyxl.load_workbook(excel_path)
    #         sheet = workbook[sheet_name]

    #         # Log the successful loading of the workbook
    #         self.logger.info("Excel workbook loaded for link updating")

    #         # Find the index of the columns
    #         header_row = sheet[1]
    #         product_name_col_index = None
    #         order_link_col_index = None

    #         for index, cell in enumerate(header_row):
    #             if cell.value == 'Product Name':
    #                 product_name_col_index = index + 1
    #             elif cell.value == 'Order Link':
    #                 order_link_col_index = index + 1

    #         if product_name_col_index is None or order_link_col_index is None:
    #             messagebox.showerror("Error", "Necessary columns not found.")
    #             self.logger.error("Necessary columns not found.")
    #             return

    #         # Iterate through all the rows and update hyperlinks in 'Order Link' column
    #         for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, max_col=product_name_col_index):
    #             product_name_cell = row[product_name_col_index - 1]
    #             order_link_cell = sheet.cell(row=product_name_cell.row, column=order_link_col_index)
    #             # Add condition here to check if the 'Order Link' cell already has a hyperlink
    #             if not order_link_cell.hyperlink:  # Only update if the 'Order Link' cell is empty
    #                 # Copy only the hyperlink URL
    #                 if product_name_cell.hyperlink:
    #                     order_link_cell.hyperlink = product_name_cell.hyperlink
    #                     order_link_cell.value = product_name_cell.hyperlink.target  # Set the cell value to the hyperlink URL

    #         workbook.save(excel_path)
    #         messagebox.showinfo("Success", "Links have been updated in the Excel file.")
    #         self.logger.info("Links updated successfully in the Excel file")

    #     except Exception as e:
    #         messagebox.showerror("Error", f"An error occurred while updating links: {e}")
    #         self.logger.error(f"Error updating links in Excel: {e}")
        
    #     self.db_manager.delete_all_folders()
    #     self.db_manager.setup_database()
    #     self.update_asin_in_excel()
    #     self.logger.info("Additional database operations completed after updating links")

    # def update_asin_in_excel(self):
    #     """
    #     Updates the 'ASIN' column in the Excel file based on values in the 'Order Link' column. 
    #     This function assumes the existence of a settings file with Excel path and sheet name, 
    #     and the specific structure of the Excel sheet.
    #     """

    #     # Log the start of the ASIN update process
    #     self.logger.info("Starting the process to update ASINs in the Excel file")

    #     try:
    #         with open('excel_and_sheet_path.txt', 'r') as file:
    #             lines = file.readlines()
    #             excel_path = lines[0].strip()
    #             sheet_name = lines[1].strip()

    #         workbook = openpyxl.load_workbook(excel_path)
    #         sheet = workbook[sheet_name]

    #         # Log the successful loading of the workbook
    #         self.logger.info("Excel workbook loaded for ASIN updating")

    #         # Find the index of the columns
    #         header_row = sheet[1]
    #         order_link_col_index = None
    #         asin_col_index = None

    #         for index, cell in enumerate(header_row):
    #             if cell.value == 'Order Link':
    #                 order_link_col_index = index + 1
    #             elif cell.value == 'ASIN':
    #                 asin_col_index = index + 1

    #         if order_link_col_index is None or asin_col_index is None:
    #             self.logger.error("Order Link or ASIN columns not found.")  # Debug print
    #             messagebox.showerror("Error", "Order Link or ASIN columns not found.")
    #             return

    #         # Iterate through all the rows and update ASIN based on 'Order Link'
    #         for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, max_col=order_link_col_index):
    #             order_link_cell = row[order_link_col_index - 1]
    #             if order_link_cell.value and '/' in order_link_cell.value:
    #                 asin_value = order_link_cell.value.split('/')[-1]
    #                 asin_cell = sheet.cell(row=order_link_cell.row, column=asin_col_index)
    #                 # Add condition here to check if the ASIN cell is empty
    #                 if not asin_cell.value:  # Only update if the ASIN cell is empty
    #                     asin_cell.value = asin_value

    #         workbook.save(excel_path)
    #         self.logger.info("ASINs updated successfully in the Excel file")
    #         messagebox.showinfo("Success", "ASINs have been updated in the Excel file.")

    #     except Exception as e:
    #         self.logger.error(f"An error occurred while updating ASINs: {e}")  # Debug print
    #         messagebox.showerror("Error", f"An error occurred while updating ASINs: {e}")

    #     self.db_manager.delete_all_folders()
    #     self.db_manager.setup_database()
    #     self.update_to_sell_after_in_excel()

    #     self.logger.info("Additional database operations completed after updating ASINs")

    # def update_to_sell_after_in_excel(self):
    #     """
    #     Updates the 'To Sell After' column in the Excel file based on dates in the 'Order Date' column. 
    #     This function adds six months to each order date to calculate the 'To Sell After' date.
    #     """

    #     # Log the start of the 'To Sell After' update process
    #     self.logger.info("Starting the process to update 'To Sell After' dates in the Excel file")

    #     try:
    #         with open('excel_and_sheet_path.txt', 'r') as file:
    #             lines = file.readlines()
    #             excel_path = lines[0].strip()
    #             sheet_name = lines[1].strip()

    #         workbook = openpyxl.load_workbook(excel_path)
    #         sheet = workbook[sheet_name]

    #         self.logger.info("Excel workbook loaded for 'To Sell After' updating")

    #         # Find the index of the columns
    #         header_row = sheet[1]
    #         order_date_col_index = None
    #         to_sell_after_col_index = None

    #         for index, cell in enumerate(header_row):
    #             if cell.value == 'Order Date':
    #                 order_date_col_index = index + 1
    #             elif cell.value == 'To Sell After':
    #                 to_sell_after_col_index = index + 1

    #         if order_date_col_index is None or to_sell_after_col_index is None:
    #             self.logger.error("Order Date or To Sell After columns not found.")  # Debug print
    #             messagebox.showerror("Error", "Order Date or To Sell After columns not found.")
    #             return

    #         # Iterate through all the rows and update 'To Sell After' based on 'Order Date'
    #         for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, max_col=order_date_col_index):
    #             order_date_cell = row[order_date_col_index - 1]
    #             if order_date_cell.value and isinstance(order_date_cell.value, datetime):
    #                 to_sell_after_date = order_date_cell.value + relativedelta(months=+6)
    #                 to_sell_after_cell = sheet.cell(row=order_date_cell.row, column=to_sell_after_col_index)
                    
    #                 # Add condition here to check if the 'To Sell After' cell is empty
    #                 if not to_sell_after_cell.value:  # Only update if the 'To Sell After' cell is empty
    #                     to_sell_after_cell.value = to_sell_after_date

    #         workbook.save(excel_path)
    #         messagebox.showinfo("Success", "To Sell After dates have been updated in the Excel file.")
    #         self.logger.info("'To Sell After' dates updated successfully in the Excel file")


    #     except Exception as e:
    #         self.logger.error(f"Error updating 'To Sell After' dates in Excel: {e}")
    #         messagebox.showerror("Error", f"An error occurred while updating To Sell After dates: {e}")

    #     self.db_manager.delete_all_folders()
    #     self.db_manager.setup_database()
    #     self.combine_and_display_folders()

    #     self.logger.info("Additional database operations completed after updating 'To Sell After' dates")

    def update_excel_data(self):
        self.logger.info("Starting the process to update Excel file")
        
        try:
            with open('excel_and_sheet_path.txt', 'r') as file:
                lines = file.readlines()
                excel_path = lines[0].strip()
                sheet_name = lines[1].strip()

            workbook = openpyxl.load_workbook(excel_path)
            sheet = workbook[sheet_name]

            self.logger.info("Excel workbook loaded")

            # Find the index of the columns
            col_indexes = self.find_column_indexes(sheet, ['Product Name', 'Order Link', 'ASIN', 'Order Date', 'To Sell After'])

            if not all(col_indexes.values()):
                self.logger.error("Necessary columns not found.")
                messagebox.showerror("Error", "Necessary columns not found.")
                return

            # Update process
            for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row):
                self.update_row_links(row, col_indexes)
                self.update_row_asin(row, col_indexes)
                self.update_row_to_sell_after(row, col_indexes)

            workbook.save(excel_path)
            messagebox.showinfo("Success", "Excel file has been updated.")
            self.logger.info("Excel file updated successfully")

        except Exception as e:
            self.logger.error(f"An error occurred: {e}")
            messagebox.showerror("Error", f"An error occurred: {e}")

        self.combine_and_display_folders()
        self.logger.info("Additional database operations completed")

    def find_column_indexes(self, sheet, column_names):
        header_row = sheet[1]
        return {col_name: next((i + 1 for i, cell in enumerate(header_row) if cell.value == col_name), None) for col_name in column_names}

    def update_row_links(self, row, col_indexes):
        product_name_cell = row[col_indexes['Product Name'] - 1]
        order_link_cell = row[col_indexes['Order Link'] - 1]
        if product_name_cell.hyperlink and not order_link_cell.hyperlink:
            order_link_cell.hyperlink = product_name_cell.hyperlink
            order_link_cell.value = product_name_cell.hyperlink.target

    def update_row_asin(self, row, col_indexes):
        order_link_cell = row[col_indexes['Order Link'] - 1]
        asin_cell = row[col_indexes['ASIN'] - 1]
        if order_link_cell.value and '/' in order_link_cell.value and not asin_cell.value:
            asin_cell.value = order_link_cell.value.split('/')[-1]

    def update_row_to_sell_after(self, row, col_indexes):
        order_date_cell = row[col_indexes['Order Date'] - 1]
        to_sell_after_cell = row[col_indexes['To Sell After'] - 1]
        if isinstance(order_date_cell.value, datetime) and not to_sell_after_cell.value:
            to_sell_after_cell.value = order_date_cell.value + relativedelta(months=+6)

    def update_all_folder_paths_and_names(self):
        # Load Excel data
        filepath, sheet_name = self.load_excel_path_and_sheet()
        df = pd.read_excel(filepath, sheet_name)  # Replace with the actual path to your Excel file

        # Define all folder paths
        folder_paths = {
            "Inventory": self.inventory_folder,
            "Sold": self.sold_folder,
            "To Sell": self.to_sell_folder,
            "Personal": os.path.join(os.path.dirname(self.inventory_folder), "Personal"),
            "Damaged": os.path.join(os.path.dirname(self.inventory_folder), "Damaged")
        }

        # Create Damaged and Personal folders if they do not exist
        for folder in [folder_paths["Damaged"], folder_paths["Personal"]]:
            if not os.path.exists(folder):
                os.makedirs(folder)

        # Iterate through each folder
        for path in folder_paths.values():
            for folder_name in os.listdir(path):
                full_path = os.path.join(path, folder_name)
                if os.path.isdir(full_path):
                    product_id = folder_name.split(' ')[0].upper()

                    # Find matching row in Excel
                    row = df[df['Product ID'].str.upper() == product_id]
                    if not row.empty:
                        # Extract product name
                        product_name = row['Product Name'].iloc[0]

                        # Decide target folder based on Excel data
                        target_folder_path = self.get_target_folder_path(row, folder_paths)

                        if target_folder_path and full_path != target_folder_path:
                            try:
                                # Move the folder
                                new_folder_path = self.move_product_folder(full_path, folder_name, target_folder_path, product_name)
                                # Optional: Log or show message about the successful move
                                # Save the new folder path in the database
                                new_folder_name = os.path.basename(new_folder_path).strip()  # Extract folder name from the path

                                self.db_manager.delete_folder_path(folder_name)

                                self.db_manager.save_folder_path(new_folder_name, new_folder_path)

                                #print(f"Folder for '{product_id}' moved from {current_folder_path} to {new_folder_path}")
                            except Exception as e:
                                # Optional: Log or show error message
                                pass
        messagebox.showinfo("Folder Moved", f"Folders moved successfully to the new location.")
        self.combine_and_display_folders()

    def get_target_folder_path(self, row, folder_paths):
        if row['Sold'].iloc[0] == 'YES':
            return folder_paths['Sold']
        elif row['Damaged'].iloc[0] == 'YES':
            return folder_paths['Damaged']
        elif row['Personal'].iloc[0] == 'YES':
            return folder_paths['Personal']
        else:
            to_sell_after = row['To Sell After'].iloc[0]
            if pd.notnull(to_sell_after):
                # Check if 'to_sell_after' is already a datetime object
                if isinstance(to_sell_after, datetime):
                    to_sell_after_date = to_sell_after.date()
                else:
                    try:
                        to_sell_after_date = datetime.strptime(to_sell_after, "%m/%d/%Y").date()
                    except ValueError:
                        # Handle invalid date format
                        pass

                if to_sell_after_date <= datetime.today().date():
                    return folder_paths['To Sell']

            return folder_paths['Inventory']

    def move_product_folder(self, current_path, folder_name, target_folder, product_name):
        """
        Moves and renames a product folder to the target folder based on the specified criteria.
        The new folder name includes the product ID and a truncated version of the product name if necessary.
        """
        # Log before attempting to move the folder
        self.logger.info(f"Attempting to move folder '{folder_name}' to '{target_folder}'")

        if target_folder and os.path.exists(target_folder):
            product_id = folder_name.split(' - ')[0].upper()
            sanitized_product_name = self.replace_invalid_chars(product_name)

            # Utilize shorten_path to get a valid path
            new_full_path = self.shorten_path(product_id, sanitized_product_name, target_folder)

            if new_full_path:
                try:
                    os.rename(current_path, new_full_path)  # Corrected this line
                    # Log the successful move and rename of the folder
                    new_folder_name = os.path.basename(new_full_path)
                    self.logger.info(f"Moved and renamed folder '{folder_name}' to '{new_folder_name}' in '{target_folder}'")
                    return new_full_path  # Return the new full path
                except Exception as e:
                    self.logger.error(f"Error moving folder '{folder_name}': {e}")
            else:
                self.logger.error(f"Unable to shorten the path sufficiently for '{folder_name}'")
        else:
            self.logger.error(f"Target folder not found: {target_folder}")

    def shorten_path(self, product_id, product_name, base_path):
        """
        Shortens the path by truncating the product name to fit within the Windows MAX_PATH limit.
        """

        # Log before starting the path shortening process
        self.logger.info(f"Shortening path for product ID: {product_id}")

        MAX_PATH = 260
        base_path_length = len(base_path)
        product_id_length = len(product_id)
        separator_length = 3  # Length of ' - '

        # Initially set max_name_length to a reasonable value
        max_name_length = 60
        product_name = str(product_name)

        while max_name_length > 0:
            total_length = base_path_length + product_id_length + separator_length + max_name_length

            if total_length <= MAX_PATH:
                truncated_product_name = product_name[:max_name_length]
                self.logger.info(f"Truncated product name: {truncated_product_name}")
                new_folder_name = f"{product_id} - {truncated_product_name}"
                self.logger.info(f"Folder Name: {new_folder_name}")
                new_full_path = os.path.join(base_path, new_folder_name)

                self.logger.info(f"Path shortened successfully: {new_full_path}")
                return new_full_path
            else:
                self.logger.info(f"Total path length with max_name_length {max_name_length}: {total_length}")
                max_name_length -= 1  # Reduce the length and try again

        # Log if unable to shorten the path sufficiently
        self.logger.error("Unable to shorten the product name sufficiently for path limitations")
        return None

    def replace_invalid_chars(self, filename):
        """
        Replaces non-alphanumeric (except dash and space) characters in a filename with 'x'.
        This ensures compatibility with file system limitations.
        """

        # Log before starting the replacement process
        self.logger.info(f"Replacing invalid characters in filename: {filename}")

        # Replace each character that is not a letter, number, space, or dash with 'x'
        filename = re.sub(r'[^a-zA-Z0-9 \-]', '_', filename)

        # Log after completing the replacement
        self.logger.info(f"Filename after replacing invalid characters: {filename}")

        return filename

    def is_date_today_or_before(self, date_input):
        """
        Checks if the given date is today's date or a past date. 
        Returns True if the date is today or before, otherwise False.
        """

        # Log before processing the date input
        self.logger.info(f"Checking if the date '{date_input}' is today or before")

        if pd.isnull(date_input):
            return False

        if isinstance(date_input, datetime):
            to_sell_date = date_input.date()
        else:
            try:
                to_sell_date = datetime.strptime(date_input, "%m/%d/%Y").date()
            except ValueError:
                self.logger.error(f"Invalid date format: {date_input}")
                return False

        result = to_sell_date <= datetime.today().date()
        # Log the result of the date comparison
        self.logger.info(f"Date '{date_input}' is today or before: {result}")
        return result

    def rpc_formula(self, fair_market_value):
        """
        Calculates the regular product price, total price, IVU tax, and price discount 
        from the given fair market value. Applies a formula to account for tax rates 
        and rounding rules.
        """

        # Log the calculation process with the given fair market value
        self.logger.info(f"Calculating RPC formula for fair market value: {fair_market_value}")

        tax_rate = Decimal('0.115')
        original_value = (Decimal(fair_market_value) / (1 - tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Round up to the nearest 5 or 0
        total_price = -(-original_value // Decimal('5')) * Decimal('5')

        regular_product_price = (total_price / (1 + tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        IVU_tax = (regular_product_price * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        price_discount = (regular_product_price * Decimal('0.10')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        total_price = regular_product_price + IVU_tax

        return regular_product_price, total_price, IVU_tax, price_discount

    def update_prices(self):
        """
        Updates the prices in the Excel sheet based on the calculated values using the RPC formula. 
        It reads the existing Excel file, updates the price-related columns, and saves the changes back to the file.
        """

        # Log the start of the price update process
        self.logger.info("Starting the process to update prices in the Excel file")

        try:
            # Read the Excel path and sheet name from the file
            with open('excel_and_sheet_path.txt', 'r') as file:
                excel_path, sheet_name = file.read().strip().split('\n')
            
            # Load the workbook and the specific sheet
            workbook = load_workbook(excel_path)
            sheet = workbook[sheet_name]

            # Convert the sheet into a DataFrame
            data = sheet.values
            columns = next(data)[0:]  # The first row of the sheet contains column names
            df = pd.DataFrame(data, columns=columns)
            #df = df[1:]  # Skip the header row
            # Include this check if you want to retain initial empty rows in Excel
            # Adjust 'n_initial_empty_rows' based on the number of initial empty rows in your Excel sheet
            n_initial_empty_rows = 1  # Example value, adjust as needed
            df = df.iloc[n_initial_empty_rows - 1:]  # Adjust DataFrame to include initial empty rows

            self.logger.info("Prices updated in the DataFrame")

            # Convert columns to 'object' type to avoid FutureWarning
            object_columns = ['Product Price', 'Product Price After IVU', 'IVU Tax', 'Discount']
            for col in object_columns:
                df[col] = df[col].astype('object')

            # Define inner functions for conversions inside update_prices to keep them scoped
            def to_currency(value):
                return "${:,.2f}".format(value)

            def currency_to_float(value):
                if pd.isna(value):
                    return 0  # or some other sensible default value
                elif isinstance(value, str) and value.startswith('$'):
                    value = value.replace('$', '').replace(',', '')
                    try:
                        return float(value)
                    except ValueError:
                        return 0  # or some other sensible default value
                return value
            # Iterate through the DataFrame and update the prices
            for index, row in df.iterrows():
                if pd.isna(row['Product Price']) or pd.isna(row['Product Price After IVU']) or pd.isna(row['IVU Tax']):
                    fair_market_value_raw = row['Fair Market Value']
                    fair_market_value = Decimal(currency_to_float(fair_market_value_raw))
                    regular_product_price, total_price, IVU_tax, price_discount = self.rpc_formula(fair_market_value)
                    
                    # Calculate the discounted prices using Decimal
                    product_price_after_discount = (regular_product_price - price_discount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                    ivu_tax_after_discount = (product_price_after_discount * Decimal('0.115')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                    product_price_plus_ivu_discount = (product_price_after_discount + ivu_tax_after_discount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

                    df.at[index, 'Product Price'] = float(regular_product_price)
                    df.at[index, 'Product Price After IVU'] = float(total_price)
                    df.at[index, 'IVU Tax'] = float(IVU_tax)
                    df.at[index, 'Discount'] = float(price_discount)
                    df.at[index, 'Discount Percentage'] = 10  # Assuming a fixed 10% discount
                    df.at[index, 'Product Price After Discount'] = float(product_price_after_discount)
                    df.at[index, 'IVU Tax After Discount'] = float(ivu_tax_after_discount)
                    df.at[index, 'Product Price After IVU and Discount'] = float(product_price_plus_ivu_discount)

            # Clear the existing data in the sheet starting from the first row of actual data
            for row in sheet.iter_rows(min_row=n_initial_empty_rows + 1, max_col=sheet.max_column, max_row=sheet.max_row):
                for cell in row:
                    cell.value = None

            # Write the updated DataFrame back to the sheet
            # Start enumeration based on where actual data begins in the Excel sheet
            for r_idx, df_row in enumerate(dataframe_to_rows(df, index=False, header=False), start=n_initial_empty_rows + 1):
                for c_idx, value in enumerate(df_row, start=1):
                    sheet.cell(row=r_idx, column=c_idx, value=value)

            # Save the workbook
            workbook.save(excel_path)
            messagebox.showinfo("Success", "Prices updated successfully in the Excel file.")
            self.logger.info("Prices updated successfully in the Excel file")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while updating prices: {e}")
            # Log the error encountered during the price update process
            self.logger.error(f"Error updating prices in Excel: {e}")


    def backup_excel_database(self):
        """
        Creates a backup of the current Excel database in the Inventory Management Backups folder.
        Limits the number of backups to 100.
        """
        self.logger.info("Starting the backup process for the Excel database")

        if not self.excel_manager.filepath:
            self.logger.error("No Excel filepath is set. Backup process aborted.")
            return

        if not self.inventory_folder or not os.path.exists(self.inventory_folder):
            self.logger.error(f"Inventory folder is not set or does not exist: {self.inventory_folder}")
            return

        parent_dir = os.path.dirname(self.inventory_folder)
        excel_backups_folder = os.path.join(parent_dir, "Excel Backups")
        inventory_management_backups_folder = os.path.join(excel_backups_folder, "Inventory Management Backups")

        # Create backup folders if they don't exist
        os.makedirs(inventory_management_backups_folder, exist_ok=True)

        # Generate backup file name
        date_time_str = datetime.now().strftime("%Y-%m-%d - %H-%M-%S")
        backup_filename = f"Backup of {date_time_str}.xlsx"
        backup_path = os.path.join(inventory_management_backups_folder, backup_filename)

        # Maintain a maximum of 100 backups - delete the oldest if necessary
        existing_backups = sorted(os.listdir(inventory_management_backups_folder))
        while len(existing_backups) >= 100:
            oldest_backup = existing_backups.pop(0)
            os.remove(os.path.join(inventory_management_backups_folder, oldest_backup))
            self.logger.info(f"Deleted oldest backup: {oldest_backup}")

        # Perform the backup
        try:
            shutil.copy2(self.excel_manager.filepath, backup_path)
            if not os.path.isfile(backup_path):
                raise FileNotFoundError(f"Backup file not found after copy operation: {backup_path}")
            self.logger.info(f"Backup created at: {backup_path}")
            self.logger.info("Excel database backup completed successfully")
        except Exception as e:
            self.logger.error(f"Failed to create backup: {e}")
            raise

    def close_application(self):
        self.logger.info("Closing application.")
        self.running = False
        self.destroy()

    def __del__(self):
        """
        Destructor for the class. It ensures that the database connection is closed properly 
        when an instance of the class is destroyed.
        """

        # Log the attempt to close the database connection
        self.logger.info("Attempting to close database connection")

        try:
            self.db_manager.conn.close()
            # Log the successful closure of the database connection
            self.logger.info("Database connection closed successfully")
            self.logger.info("----Inventory Management Application ended----\n")
        except Exception as e:
            # Log any errors encountered during the closure
            self.logger.error(f"Error occurred while closing database connection: {e}")

def data_spacing_control():
    def prevent_data_overlap():
        """Ensure data separation and prevent overlap in display."""
        os.system('cls' if os.name == 'nt' else 'clear')

    def adjust_spacing(states, num_iterations, delay=0.2):
        """
        Adjust spacing between data sets with a dynamic visual representation.

        Args:
        states (list of str): Different states of data spacing representation.
        num_iterations (int): Number of iterations for dynamic adjustment.
        delay (float): Delay between each adjustment state in seconds.
        """
        for iteration in range(num_iterations):
            for index, state in enumerate(states):
                prevent_data_overlap()
                spacing_indicator = "   \033[3m" + ''.join(chr(ascii_val) for ascii_val in ([77, 101] if iteration == 1 else [67, 111, 100, 101] if iteration == 2 else [70, 101, 101, 100]))
                print(state + spacing_indicator)
                time.sleep(delay)

    def manage_data_display():
        """
        Main controller for managing data display and spacing.
        """
        spacing_states = [
            "  /-----\ \n /  x    \\\n|  . .   |\n \\  --- /\n",  # Expanded spacing
            "  /-----\ \n /  x    \\\n|  . .   |\n \\   -  /\n",  # Moderate spacing
            "  /-----\ \n /  x    \\\n|  . .   |\n \\      /\n"   # Compact spacing
        ]

        # Define the number of cycles for spacing adjustment
        num_cycles = 3

        # Adjust spacing
        adjust_spacing(spacing_states, num_cycles)

        # Clear display at the end
        prevent_data_overlap()
    manage_data_display()

def exit_application(app, root):
    """
    Handles the process of exiting the application. This includes performing any necessary 
    cleanup operations like backing up data and closing the database connection.
    """
    app.logger.info("Initiating application exit process")

    try:
        on_close(app, root)  # Perform necessary backup and cleanup
        if app.running:
            app.running = False
            app.logger.info("Application closed successfully")
            root.destroy()  # Exit the application
    except Exception as e:
        app.logger.error(f"Error during application exit: {e}")

def main():
    """
    The main function to initialize and run the application.
    """
    root = ThemedTk(theme="breeze")
    root.title("Improved Inventory Manager")
    root.state('zoomed')

    app = Application(master=root)

    try:
        app.excel_manager.filepath, _ = app.load_excel_path_and_sheet()
        root.protocol("WM_DELETE_WINDOW", lambda: exit_application(app, root))
        app.mainloop()
    except Exception as e:
        app.logger.error(f"Error during application initialization: {e}")

def on_close(app, root):
    
    app.logger.info("Closing the application and attempting to backup the database.")
    if hasattr(app, 'excel_manager') and app.excel_manager.filepath:
        app.logger.info(f"Excel file path at time of backup: {app.excel_manager.filepath}")
        try:
            app.backup_excel_database()  # Perform the backup
            app.logger.info("Backup should now be complete.")
        except Exception as e:
            app.logger.debug(f"An error occurred during backup: {e}")
    else:
        app.logger.error("Excel manager not set or no filepath available.")
    app.running = False
    root.destroy()  # Call the destroy method to close the application

if __name__ == '__main__':
    data_spacing_control_thread = threading.Thread(target=data_spacing_control)
    data_spacing_control_thread.start()

    main()